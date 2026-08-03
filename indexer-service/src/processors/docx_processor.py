"""
DOCX Document Processor

Extracts text from Microsoft Word documents using python-docx.
"""
import io
from typing import Optional

from docx import Document

from core.logging import get_logger

from .base import DocumentProcessor, ProcessedDocument, ProcessingException

logger = get_logger(__name__)


class DOCXProcessor(DocumentProcessor):
    """
    DOCX document processor.

    Uses python-docx to extract text from Microsoft Word documents.
    """

    async def can_process(self, file_name: str, content_type: Optional[str] = None) -> bool:
        """Check if file is a DOCX"""
        if content_type and "wordprocessingml" in content_type.lower():
            return True

        return file_name.lower().endswith((".docx", ".doc"))

    async def process(self, content: bytes, file_name: str) -> ProcessedDocument:
        """Process DOCX and extract text"""
        try:
            logger.info("Processing DOCX document", file_name=file_name)

            text = await self.extract_text(content)

            if not text or len(text.strip()) < 10:
                raise ProcessingException(
                    message="Could not extract text from DOCX",
                    file_name=file_name,
                )

            # Count paragraphs as rough page estimate
            paragraph_count = len(text.split("\n\n"))
            page_estimate = max(1, paragraph_count // 5)

            logger.info(
                "DOCX processed successfully",
                file_name=file_name,
                text_length=len(text),
                page_estimate=page_estimate,
            )

            return ProcessedDocument(
                text=text,
                metadata={
                    "file_name": file_name,
                    "file_type": "docx",
                    "processor": "docx_processor",
                },
                page_count=page_estimate,
            )

        except Exception as e:
            logger.error("Failed to process DOCX", error_msg=e, file_name=file_name)
            raise ProcessingException(
                message=f"Failed to process DOCX: {str(e)}",
                file_name=file_name,
                cause=e,
            )

    async def extract_text(self, content: bytes) -> str:
        """Extract plain text from DOCX"""
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)

            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.warning("DOCX extraction failed", error_msg=e)
            raise

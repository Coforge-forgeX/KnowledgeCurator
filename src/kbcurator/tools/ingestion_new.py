import os
from dotenv import load_dotenv
import time
import traceback
from lightrag.utils import EmbeddingFunc
from lightrag import LightRAG,QueryParam

from lightrag.llm.bedrock import bedrock_complete, bedrock_embed

from lightrag.kg import STORAGES, STORAGE_IMPLEMENTATIONS, verify_storage_implementation

from common_adapters.ai.unified import UnifiedAIAdapter, AzureConfig, AWSConfig
from common_adapters.lightrag.neptune.neptune_impl import NeptuneGraphStorage
from common_adapters.lightrag.neptune.neptune_to_neo4j_converter import NeptuneToNeo4jConverter
from common_adapters.doc_extract import DocReader
from common_adapters.storage.exceptions import StorageError, NotFoundError
from kbcurator.server import storage_config

import inspect

from lightrag.kg.shared_storage import initialize_share_data, initialize_pipeline_status
import aiohttp
from kbcurator.server.server import mcp
import psycopg2
from azure.storage.blob import BlobServiceClient
from docx import Document
import io
import numpy as np
from typing import Optional, List, Any
from lightrag import QueryParam
from fastmcp import Context
import asyncio
import json
import zipfile
import tempfile
from azure.storage.blob import BlobServiceClient, ContentSettings
from azure.storage.blob import generate_blob_sas, BlobSasPermissions
from datetime import datetime, timedelta
import base64
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
from crawl4ai import AsyncWebCrawler
from crawl4ai.async_configs import BrowserConfig, CrawlerRunConfig, CacheMode, DefaultMarkdownGenerator
from kbcurator.utils.azurecustomllm import AzureCustomLLM
from kbcurator.utils.access_validation import validate_user_workspace_access
from kbcurator.utils.request_context import request_var
from kbcurator.tools.user_management_system import Session,UserMap
 
load_dotenv(os.path.abspath(os.path.join(os.getcwd(),'.env')))
 
azure_llm_api_key = os.getenv('AZURE_OPENAI_LLM_MODEL_API_KEY')
azure_llm_api_base = os.getenv('AZURE_OPENAI_LLM_MODEL_API_BASE')
azure_llm_api_version = os.getenv('AZURE_OPENAI_LLM_MODEL_API_VERSION')
azure_llm_deployement_name = os.getenv('AZURE_OPENAI_LLM_MODEL_LLM_MODEL')
 
azure_embedding_api_key = os.getenv('AZURE_OPENAI_EMBEDDING_MODEL_API_KEY')
azure_embedding_api_base = os.getenv('AZURE_OPENAI_EMBEDDING_MODEL_API_BASE')
azure_embedding_api_version = os.getenv('AZURE_OPENAI_EMBEDDING_MODEL_API_VERSION')
azure_embedding_deployement_name = os.getenv('AZURE_OPENAI_EMBEDDING_MODEL_EMBEDDING_MODEL')
 
 
os.environ["NEO4J_URI"] = os.getenv("NEO4J_DATABASE_NEO4J_BOLT_URI", "bolt://localhost:7687") or "bolt://localhost:7687"
os.environ["NEO4J_USERNAME"] = os.getenv("NEO4J_DATABASE_NEO4J_USER") or ""
os.environ["NEO4J_PASSWORD"] = os.getenv("NEO4J_DATABASE_NEO4J_PASSWORD") or ""
 
# os.environ["POSTGRES_HOST"] = os.getenv("POSTGRESQL_DATABASE_HOST") or ""
os.environ["POSTGRES_HOST"] = os.getenv("LIGHTRAG_POSTGRESQL_DATABASE_HOST") or ""
os.environ["POSTGRES_USER"] = os.getenv("LIGHTRAG_POSTGRESQL_DATABASE_USER") or ""
os.environ["POSTGRES_PASSWORD"] = os.getenv("LIGHTRAG_POSTGRESQL_DATABASE_PASSWORD") or ""
os.environ["POSTGRES_DATABASE"] = os.getenv("LIGHTRAG_POSTGRESQL_DATABASE_DATABASE") or ""
 
embedding_dim = int(os.getenv("OLLAMA_MODEL_EMBEDDING_MODEL_DIMS", "3072"))
max_token_size = int(os.getenv("OLLAMA_MODEL_EMBEDDING_MODEL_MAX_TOKENS", "8192"))
base_url = os.getenv("OLLAMA_MODEL_BASE_URL")
 
DEFAULT_TOP_K = 5
MAX_TOP_K = 20
STREAMING = True

# 1) Choose provider via ENV (no code changes when switching)
PROVIDER = os.getenv("MODEL_PROVIDER", "azure")  # "azure" or "bedrock"

adapter = UnifiedAIAdapter(
        provider=PROVIDER
    )

async def llm_model_func(prompt, system_prompt=None, history_messages=[], **kwargs) -> str:
    # kwargs like temperature/top_p/max_tokens/timeouts pass through
    return await adapter.acomplete(
        prompt=prompt, 
        messages=history_messages, 
        )

async def embedding_func(texts: list[str]) -> np.ndarray:
    return await adapter.embed(
        texts=texts
        )
 
async def initialize_rag(domain: Optional[str] = None, kb_name: Optional[str] = None) -> LightRAG:
    data_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'data'))
    lightrag_database = ''.join(char for char in f"{domain}{kb_name}" if char.isalpha())
    os.environ['NEO4J_DATABASE'] = lightrag_database
    print(''.join(char for char in f"{domain}{kb_name}" if char.isalpha()))

    STORAGES["NeptuneGraphStorage"] = "common_adapters.lightrag.neptune.neptune_impl"
    
    impls = STORAGE_IMPLEMENTATIONS.get("GRAPH_STORAGE", {}).get("implementations")
    if impls is not None and "NeptuneGraphStorage" not in impls:
        impls.append("NeptuneGraphStorage")

    verify_storage_implementation(
        storage_type="GRAPH_STORAGE",
        storage_name="NeptuneGraphStorage"
        )
    
    rag = LightRAG(
            working_dir=data_dir,
            llm_model_func=bedrock_complete,
            llm_model_name=os.getenv("BEDROCK_LLM_MODEL_ID"),
            embedding_func=EmbeddingFunc(
                embedding_dim=embedding_dim,
                max_token_size=max_token_size,
                model_name=os.getenv("BEDROCK_EMBED_MODEL_ID"),
                func=bedrock_embed.func
            ),
            # graph_storage="Neo4JStorage",
            graph_storage="NeptuneGraphStorage",
            workspace = lightrag_database,
            vector_storage="PGVectorStorage",
            chunk_token_size=1000,
            chunk_overlap_token_size=200,
            addon_params={
                        # delivered to storage.global_config["addon_params"]
                        "NEPTUNE_GATEWAY_URL": os.getenv("NEPTUNE_ENDPOINT"),
                    },
        )
    await rag.initialize_storages()
    initialize_share_data()
    await initialize_pipeline_status()
    return rag
 
# @mcp.tool()
# async def query_rag(
#     domain: Optional[str] = None, 
#     kb_name: Optional[str] = None, 
#     knowledge_bases: Optional[list[str]] = None, 
#     question: Optional[str] = None, 
#     user_prompt: Optional[str] = None, 
#     history: Optional[list] = None,
#     mode: str = 'mix'
# ) -> dict:
#     """
#     Query the RAG system with a question and optional user prompt.
#     If knowledge_bases is provided, query each and aggregate results.
#     """
#     if history is None:
#         history = []
#     try:
#         user_prompt = f"""---Role---

# You are a helpful assistant responding to user queries about Knowledge Graph and Document Chunks provided in JSON format below.

# ---Goal---

# Generate a concise, accurate response based on the provided Knowledge Base. Follow all Response Rules strictly. Use both the conversation history and the current query to guide your response. Do not include any information not present in the Knowledge Base or conversation history.

# When handling relationships with timestamps:
# 1. Each relationship has a "created_at" timestamp indicating when we acquired this knowledge.
# 2. When encountering conflicting relationships, consider both semantic content and timestamp.
# 3. Do not automatically prefer the most recent relationship—use contextual judgment.
# 4. For time-specific queries, prioritize temporal information in the content before considering timestamps.

# ---Query to be answered---
# {question}

# ---Conversation History---
# {history}

# ---Response Rules---

# - **Format**: Use multiple paragraphs with markdown formatting and section headings.
# - **Language**: Respond in the same language as the user's question.
# - **Emphasis**: Highlight all referenced information using **bold text**.
# - **Continuity**: Maintain logical flow with the conversation history.

# - **Inline Citations**:
#     - Cite the source **immediately after** the referenced information using square brackets (e.g., [1], [2]).
#     - Every time a source is used, it must be cited—even if it has been cited earlier.

# - **Reference Mapping**:
#     - Assign each source file a unique reference number starting from 1, in the order of **first appearance** in the main answer.
#     - Maintain a mapping between source file names and their assigned reference numbers.
#     - Use this mapping consistently throughout the answer.

# - **References Section**:
#     - Include a "References" section at the end listing only the **file names** (not full paths) that were cited in the main answer.
#     - List them in ascending order of their citation number.
#     - Each file name should appear **only once**.
#     - Do not include any file name that was not cited in the main answer.
#     - Ensure that every reference number used in the main answer appears exactly once in the "References" section.

# - **Integrity**:
#     - If the answer is unknown, say so clearly.
#     - Do not fabricate information or include anything not present in the Knowledge Base or conversation history."""

#         print(
#             "Domain:", domain, 
#             "kb_name:", kb_name, 
#             "knowledge_bases:", knowledge_bases,
#             "question:", question, 
#             "user_prompt:", user_prompt[:10], 
#             "history:", history[:10]
#         )
#         if knowledge_bases:
#             llm_summarize = AzureCustomLLM()
#             results = {}
#             kb_graph_refs = []
#             async def query_single_kb(kb):
#                 try:
#                     # Initialize rag for each KB
#                     rag = await initialize_rag(domain=domain, kb_name=kb_name+kb)
#                     response = await rag.aquery(
#                         question if question else "",
#                         param=QueryParam(
#                             mode=mode,
#                             top_k=2,
#                             conversation_history=history,
#                             user_prompt=user_prompt,
#                             stream=False
#                         )
#                     )
#                     return (kb, response, None)
#                 except Exception as e:
#                     return (kb, None, str(e))

#             tasks = [query_single_kb(kb) for kb in knowledge_bases]
#             task_results = await asyncio.gather(*tasks)
#             for kb, response, error in task_results:
#                 if error:
#                     results[kb] = {"error": error}
#                 else:
#                     results[kb] = response
#                     kb_graph_refs.append(f"Knowledge Base: {kb}")

#             # Summarize the results using AzureCustomLLM
#             summary_prompt = (
#             """
#                 ### ---Role---

#                 You are a helpful assistant responding to user queries about Knowledge Graphs and Document Chunks provided in JSON format below.

#                 ### ---Goal---

#                 Generate a concise, accurate response based on the provided Knowledge Bases/Graphs. For each piece of information, clearly indicate which knowledge graph and file path it came from using inline citations (e.g., [KB1: /path/to/file], [KB2: /another/file]). Use both the conversation history and the current query to guide your response. Do not include any information not present in the Knowledge Base or conversation history.

#                 When handling relationships with timestamps:
#                 1. Each relationship has a "created_at" timestamp indicating when we acquired this knowledge.
#                 2. When encountering conflicting relationships, consider both semantic content and timestamp.
#                 3. Do not automatically prefer the most recent relationship—use contextual judgment.
#                 4. For time-specific queries, prioritize temporal information in the content before considering timestamps.

#                 ### ---Query to be answered---
#                 {question}

#                 ### ---Conversation History---
#                 {history}

#                 ### ---Response Rules---

#                 - **Format**: Use multiple paragraphs with markdown formatting and section headings.
#                 - **Language**: Respond in the same language as the user's question.
#                 - **Emphasis**: Highlight all referenced information using **bold text**.
#                 - **Continuity**: Maintain logical flow with the conversation history.

#                 - **Inline Citations**:
#                     - Cite the source knowledge graph and file path **immediately after** the referenced information using square brackets (e.g., [KB1: /path/to/file], [KB2: /another/file]).
#                     - Every time a source is used, it must be cited—even if it has been cited earlier.

#                 - **Reference Mapping**:
#                     - Assign each knowledge graph a unique reference number starting from 1, in the order of **first appearance** in the main answer.
#                     - Maintain a mapping between knowledge graph names and their assigned reference numbers.
#                     - Use this mapping consistently throughout the answer.

#                 - **References Section**:
#                     - Include a "References" section at the end listing only the **knowledge graph names** and their associated file paths (not full paths if not available) that were cited in the main answer.
#                     - List them in ascending order of their citation number.
#                     - Each knowledge graph name and file path pair should appear **only once**.
#                     - Do not include any knowledge graph name or file path that was not cited in the main answer.
#                     - Ensure that every reference number used in the main answer appears exactly once in the "References" section.

#                 - **Integrity**:
#                     - If the answer is unknown, say so clearly.
#                     - Do not fabricate information or include anything not present in the Knowledge Base or conversation history.

#                 ### ---Knowledge Graphs Used---
#                 {list of knowledge graphs, e.g., KB1: Banking, KB2: Insurance, ...}
#             """
#             )
#             for kb, resp in results.items():
#                 summary_prompt += f"### Knowledge Base: {kb}\n"
#                 if isinstance(resp, dict) and 'error' in resp:
#                     summary_prompt += f"Error: {resp['error']}\n"
#                 else:
#                     summary_prompt += f"Response: {str(resp)}\n"
#             summary_prompt += "\n---\nReferences:\n"
#             for i, kb in enumerate(results.keys(), 1):
#                 summary_prompt += f"[{i}] {kb}\n"
#             summary = llm_summarize._call(
#                 input=summary_prompt
#             )
#             return {"LightRAG": summary}
#         else:
#             rag = await initialize_rag(domain=domain, kb_name=kb_name)
#             response = await rag.aquery(
#                 question if question else "",
#                 param=QueryParam(
#                     mode=mode, 
#                     top_k=2, 
#                     conversation_history=history, 
#                     user_prompt=user_prompt, 
#                     stream=False
#                     )
#                 )
#             return {"LightRAG": response}
#     except Exception as e:
#             return {"error": str(e)}

@mcp.tool()
async def query_rag(
    domain: Optional[str] = None, 
    kb_name: Optional[str] = None, 
    knowledge_bases: Optional[list[str]] = None, 
    question: Optional[str] = None, 
    user_prompt: Optional[str] = None, 
    history: Optional[list] = None,
    mode: str = 'mix'
) -> dict:
    """
    Query the RAG system with a question and optional user prompt.
    If knowledge_bases is provided, query each and aggregate results.
    """

    # ---------------------------
    # Adapter Initialization
    # ---------------------------
    PROVIDER = os.getenv("MODEL_PROVIDER", "azure")  # "azure" or "bedrock"
    adapter = UnifiedAIAdapter(provider=PROVIDER)

    async def llm_generate(prompt: str, history_msgs: list = []):
        return await adapter.acomplete(
            prompt=prompt,
            messages=history_msgs
        )

    async def embed_texts(texts: list[str]):
        return await adapter.embed(texts=texts)

    # -----------------------------------
    # History sanitation (same as ingestion_new)
    # -----------------------------------
    def _to_text_block(value: Any) -> str:
        if isinstance(value, str):
            return value
        if isinstance(value, dict):
            return json.dumps(value, ensure_ascii=False)
        if isinstance(value, list):
            return "\n".join(_to_text_block(v) for v in value)
        return str(value)

    _ROLE_MAP = {
        "system": "system",
        "sys": "system",
        "developer": "system",
        "user": "user",
        "human": "user",
        "assistant": "assistant",
        "ai": "assistant",
        "bot": "assistant",
        "model": "assistant",
    }
    _ALLOWED = {"system", "user", "assistant"}

    def sanitize_history(history_list):
        if not history_list:
            return []
        out = []
        for turn in history_list:
            if not isinstance(turn, dict):
                out.append({"role": "user", "content": _to_text_block(turn)})
                continue
            raw_role = (turn.get("role") or turn.get("sender") or "user").lower()
            role = _ROLE_MAP.get(raw_role)
            if role not in _ALLOWED:
                continue
            content = turn.get("content") or turn.get("text") or turn.get("message")
            if content is None:
                continue
            out.append({"role": role, "content": _to_text_block(content)})
        return out

    # sanitize
    if history is None:
        history = []
    sanitized_history = sanitize_history(history)

    # ---------------------------
    # Build system prompt
    # ---------------------------
    system_prompt = f"""
---Role---

You are a helpful assistant responding to user queries about Knowledge Graph and Document Chunks provided in JSON format below.

---Goal---

Generate a concise, accurate response based on the provided Knowledge Base. Follow all Response Rules strictly. Use both the conversation history and the current query to guide your response. Do not include any information not present in the Knowledge Base or conversation history.

When handling relationships with timestamps:
1. Each relationship has a "created_at" timestamp indicating when we acquired this knowledge.
2. When encountering conflicting relationships, consider both semantic content and timestamp.
3. Do not automatically prefer the most recent relationship—use contextual judgment.
4. For time-specific queries, prioritize temporal information in the content before considering timestamps.

---Query to be answered---
{question}

---Conversation History---
{history}

---Response Rules---
- **Format**: Use multiple paragraphs with markdown formatting and section headings.
- **Language**: Respond in the same language as the user's question.
- **Emphasis**: Highlight all referenced information using **bold text**.
- **Continuity**: Maintain logical flow with the conversation history.

- **Inline Citations**:
    - Cite the source **immediately after** the referenced information using square brackets (e.g., [1], [2]).
    - Every time a source is used, it must be cited—even if it has been cited earlier.

- **Reference Mapping**:
    - Assign each source file a unique reference number starting from 1, in the order of **first appearance** in the main answer.
    - Maintain this mapping consistently.

- **References Section**:
    - Include a "References" section listing only the file names cited in the answer.
    - Each number appears exactly once.

- **Integrity**:
    - If the answer is unknown, say so.
    - Do not fabricate information.
"""

    # -------------------------
    # Provider transforms
    # -------------------------
    def azure_transform():
        out = [{"role": "system", "content": system_prompt}]
        for m in sanitized_history:
            out.append({"role": m["role"], "content": m["content"]})
        return out

    def bedrock_transform():
        # same logic as ingestion_new
        system_blocks = [{"text": system_prompt}]
        messages = []
        for m in sanitized_history:
            if m["role"] == "system":
                system_blocks.append({"text": m["content"]})
                continue
            messages.append({
                "role": m["role"],
                "content": [{"text": m["content"]}]
            })
        return messages, system_blocks

    prov = PROVIDER.lower()
    if prov in ("azure", "azure_openai", "azure-openai"):
        transformed_history = azure_transform()
    else:
        transformed_history = bedrock_transform()

    # -------------------------
    # Main logic (unchanged)
    # -------------------------
    print(
        "Domain:", domain, 
        "kb_name:", kb_name, 
        "knowledge_bases:", knowledge_bases,
        "question:", question, 
        "history:", history[:10]
    )

    try:
        if knowledge_bases:
            results = {}

            async def query_one(kb):
                try:
                    rag = await initialize_rag(domain=domain, kb_name=kb_name + kb)
                    response = await rag.aquery(
                        question if question else "",
                        param=QueryParam(
                            mode=mode,
                            top_k=2,
                            conversation_history=sanitized_history,
                            user_prompt=system_prompt,
                            stream=False
                        )
                    )
                    return kb, response, None
                except Exception as e:
                    return kb, None, str(e)

            task_results = await asyncio.gather(*[query_one(kb) for kb in knowledge_bases])

            # attach responses
            for kb, resp, err in task_results:
                if err:
                    results[kb] = {"error": err}
                else:
                    results[kb] = resp

            # -------------------------
            # Summarization using UnifiedAIAdapter
            # -------------------------
            summary_prompt = system_prompt

            for kb, resp in results.items():
                summary_prompt += f"\n--- KB: {kb} ---\n{str(resp)}\n"

            summary_prompt += "\nProvide a structured merged answer following all rules."

            summary = await llm_generate(summary_prompt)

            return {"LightRAG": summary}

        # -----------------------------------
        # Single KB path (unchanged, except adapter)
        # -----------------------------------
        rag = await initialize_rag(domain=domain, kb_name=kb_name)

        response = await rag.aquery(
            question if question else "",
            param=QueryParam(
                mode=mode,
                top_k=2,
                conversation_history=sanitized_history,
                user_prompt=system_prompt,
                stream=False
            )
        )
        return {"LightRAG": response}

    except Exception as e:
        print(f"Query rag error: {e}")
        return {"error": str(e)}
    
async def index_file(ctx: Context, container_client, domain, kb_name, file_path):
    print(f"Indexing file started: {file_path}")
    try:
        blob_client = container_client.get_blob_client(file_path)
        blob_data = blob_client.download_blob().readall()
        # Pass Data to your indexing function here
        resp_rag = await lightrag_indexing_tool_new(Context, blob_data, domain, kb_name, file_path)
        print(f"Response from RAG Indexing file completed: {resp_rag}")
    except Exception as e:
        return f"error: {e}"
 
async def index_sub_industry(ctx: Context, container_name, domain, kb_name, task_id=None):
    try:
        industry = domain
        sub_industry = kb_name
        max_concurrent = 3
        file_status = []
        prefix = f"{industry}/{sub_industry}/"
 
        connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        container_client = blob_service_client.get_container_client(container_name or "")
       
 
        # List all files in the subfolder
        blob_list = container_client.list_blobs(name_starts_with=prefix)
        semaphore = asyncio.Semaphore(max_concurrent)
        async def sem_index_file(file_path):
            async with semaphore:
                await index_file(Context, container_client, industry, sub_industry, file_path)
       
        tasks = [sem_index_file(blob.name) for blob in blob_list]
        await asyncio.gather(*tasks)
    except Exception as e:
        return f"Error indexing {industry}/{sub_industry}: {e}"
   
@mcp.tool()  
def start_workspace_indexing(ctx: Context, container_name: Optional[str] = None, domain: Optional[str] = None, kb_name: Optional[List[str]] = None, file_path: Optional[str] = None):
    try:
        industry = domain
        sub_industries = kb_name
        for sub_industry in sub_industries:
            task_id = f"{industry}_{sub_industry}_{int(time.time())}"
            print(f"Starting indexing for {industry}/{sub_industry} with task_id {task_id}")
            task = asyncio.create_task(index_sub_industry(Context, container_name, industry, sub_industry, task_id))
        return {"status": "success"}
    except Exception as e:
        return {"error": str(e)}
   
 
@mcp.tool()
async def index_uploaded_files(ctx: Context, container_name, domain, kb_names, file_names):
    """
    For each kb_name and file_name, generate path domain/kb_name/file_name and index it.
    """
    try:
        if not file_names:
            return {"error": "Files are required for indexing."}
       
        connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        container_client = blob_service_client.get_container_client(container_name or "")
 
        tasks = []
        for kb in kb_names:
            for file_name in file_names:
                file_path = f"{domain}/{kb}/{file_name}"
                print(f"Scheduling indexing for: {file_path}")
                tasks.append(index_file(ctx, container_client, domain, kb, file_path))
       
        results = await asyncio.gather(*tasks)
        return {"status": "success", "tasks": results}
    except Exception as e:
        return {"error": str(e)}
   
async def lightrag_indexing_tool_new(ctx: Context, blob_data, domain, kb_name, file_path) -> dict:
    """
    Index all .txt files in the specified domain/KB directory.
    """
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        ext = file_path.lower().split('.')[-1]
        content = None
       
        if ext == 'txt':
            content = blob_data.decode('utf-8', errors='ignore')
        # elif ext == 'pdf':
        #     import fitz  # PyMuPDF
        #     from paddleocr import PaddleOCR
        #     from PIL import Image
        #     import numpy as np
        #     ocr = PaddleOCR(use_angle_cls=True, lang='en')
        #     content = ""
        #     with fitz.open(stream=blob_data, filetype="pdf") as doc:
        #         all_text = []
        #         for page_num in range(len(doc)):
        #             page = doc.load_page(page_num)
        #             # Use get_pixmap for PyMuPDF >= 1.18, fallback to getPixmap for older
        #             if hasattr(page, 'get_pixmap'):
        #                 pix = page.get_pixmap()
        #             elif hasattr(page, 'getPixmap'):
        #                 pix = page.getPixmap()
        #             else:
        #                 raise RuntimeError('No get_pixmap or getPixmap method found on PyMuPDF Page object')
        #             img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        #             img_np = np.array(img)
        #             result = ocr.predict(img_np)
        #             page_text = " ".join(result[0]['rec_texts'])
        #             print(f"OCR result for {file_path} page {page_num}:", result[0]['rec_texts'])
        #             all_text.append(page_text)
        #         content = "\n".join(all_text)
        #     print(content)
        elif ext in ['docx']:
            # Ensure blob data represents raw DOCX/DOC bytes (not base64 text)
            def _ensure_doc_bytes(data: bytes | str) -> bytes:
                try:
                    if isinstance(data, (bytes, bytearray)):
                        sample = bytes(data[:32])
                        # If the start looks ASCII, try strict base64 decode
                        is_ascii_like = all(32 <= b <= 126 or b in (9, 10, 13) for b in sample)
                        if is_ascii_like:
                            try:
                                decoded = base64.b64decode(data, validate=True)
                                # DOCX files are ZIPs starting with PK\x03\x04
                                if decoded.startswith(b"PK\x03\x04"):
                                    return decoded
                            except Exception:
                                pass
                        return bytes(data)
                    elif isinstance(data, str):
                        # Handle data URLs or raw base64 strings
                        if data.startswith("data:"):
                            try:
                                b64_part = data.split("base64,", 1)[-1]
                                return base64.b64decode(b64_part, validate=True)
                            except Exception:
                                return data.encode("utf-8")
                        try:
                            return base64.b64decode(data, validate=True)
                        except Exception:
                            return data.encode("utf-8")
                    else:
                        return bytes(data)
                except Exception:
                    return bytes(data) if isinstance(data, (bytes, bytearray)) else str(data).encode("utf-8")

            doc_bytes = _ensure_doc_bytes(blob_data)
            doc = Document(io.BytesIO(doc_bytes))
            content = "\n".join([p.text for p in doc.paragraphs])
        elif ext == 'doc':
            # Attempt Windows COM-based conversion (.doc -> .docx) if Word is available
            def _convert_doc_to_docx_bytes(doc_bytes: bytes) -> bytes | None:
                try:
                    import win32com.client as win32
                    import pythoncom
                    # Write input .doc to a temp file
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as f_in:
                        f_in.write(doc_bytes)
                        in_path = f_in.name
                    out_path = in_path[:-4] + '.docx'
                    try:
                        pythoncom.CoInitialize()
                        word = win32.Dispatch('Word.Application')
                        word.Visible = False
                        doc = word.Documents.Open(in_path)
                        wdFormatXMLDocument = 12
                        doc.SaveAs(out_path, FileFormat=wdFormatXMLDocument)
                        doc.Close(False)
                        word.Quit()
                        pythoncom.CoUninitialize()
                        with open(out_path, 'rb') as f_out:
                            return f_out.read()
                    finally:
                        # Cleanup temp files
                        try:
                            import os
                            if os.path.exists(in_path):
                                os.remove(in_path)
                            if os.path.exists(out_path):
                                os.remove(out_path)
                        except Exception:
                            pass
                except Exception:
                    return None

            # Detect typical .doc OLE header
            hdr = bytes(blob_data[:8]) if isinstance(blob_data, (bytes, bytearray)) else b''
            if hdr.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1") or hdr.startswith(b"0M8R4KG"):
                converted = _convert_doc_to_docx_bytes(blob_data if isinstance(blob_data, (bytes, bytearray)) else blob_data.encode('utf-8'))
                if converted:
                    doc = Document(io.BytesIO(converted))
                    content = "\n".join([p.text for p in doc.paragraphs])
                else:
                    # Fallback: extract readable ASCII sequences from .doc bytes
                    try:
                        import re
                        raw = blob_data if isinstance(blob_data, (bytes, bytearray)) else str(blob_data).encode('utf-8', errors='ignore')
                        # Replace non-printable with spaces, keep basic punctuation
                        text = raw.decode('latin-1', errors='ignore')
                        # Find sequences of 5+ printable chars
                        blocks = re.findall(r"[\x20-\x7E]{5,}", text)
                        content = "\n".join(blocks)
                        if not content or len(content.strip()) < 20:
                            return {"error": "Legacy .doc detected and conversion unavailable. Could not extract sufficient text. Please convert to .docx and re-upload."}
                    except Exception:
                        return {"error": "Legacy .doc detected and conversion unavailable. Could not extract text reliably. Please convert to .docx and re-upload."}
            else:
                return {"error": "Unknown .doc content format. Only .docx is supported for indexing."}
        else:
            return {"error": f"Unsupported file type: {ext}"}
        def chunk_text(text, chunk_size=2000):
            # Simple chunking by character count
            return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
 
        chunks = chunk_text(content, chunk_size=2000)
        for idx, chunk in enumerate(chunks):
            await rag.ainsert(input=chunk, file_paths=[file_path])
            await ctx.debug(f"Progress: {idx+1}/{len(chunks)}")
        return {"status": "success", "file": file_path, "chunks": len(chunks)}
    except Exception as e:
        return {"error": str(e)}
   
def uploaded_by_username(uploaded_by_value: str | None) -> str | None:
    """
    If uploaded_by is a numeric user_id, resolve it to 'First Last' from users table.
    Otherwise return as-is.
    """
    try:
        if uploaded_by_value is None:
            return None
        if str(uploaded_by_value).isdigit():
            conn = psycopg2.connect(
                host=os.environ["POSTGRES_HOST"],
                user=os.environ["POSTGRES_USER"],
                password=os.environ["POSTGRES_PASSWORD"],
                dbname=os.environ["POSTGRES_DATABASE"]
            )
            cur = conn.cursor()
            cur.execute(
                "SELECT COALESCE(first_name,'') AS fn, COALESCE(last_name,'') AS ln FROM users WHERE user_id = %s",
                (int(uploaded_by_value),)
            )
            row = cur.fetchone()
            cur.close()
            conn.close()
            if row:
                fn, ln = row
                full_name = (fn + " " + ln).strip()
                return full_name if full_name else str(uploaded_by_value)
        return uploaded_by_value
    except Exception as e:
        print(f"Warning: could not resolve uploaded_by to name: {e}")
        return uploaded_by_value

def create_file_task_record(container_name, upload_path, domain, kb_name, file_path, workspace_id, status="uploading", file_size=None, uploaded_by=None):
    try:
        # Resolve uploaded_by (user_id -> full name)
        uploaded_by_resolved = uploaded_by_username(uploaded_by)

        conn = psycopg2.connect(
            host=os.environ["POSTGRES_HOST"],
            user=os.environ["POSTGRES_USER"],
            password=os.environ["POSTGRES_PASSWORD"],
            dbname=os.environ["POSTGRES_DATABASE"]
        )
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO file_tasks
            (container_name, upload_path, domain, kb_name, file_path, workspace_id, status, file_size, uploaded_by, created_at, updated_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(), NOW())
            RETURNING id;
        """, (container_name, upload_path, domain, kb_name, file_path, workspace_id, status, file_size, uploaded_by_resolved))
        task_id = cur.fetchone()[0]
        conn.commit()
        cur.close()
        conn.close()
        return task_id
    except Exception as e:
        print(f"Error creating file_tasks record: {e}")
        return None
 
def update_file_task_status(task_id, status):
    try:
        conn = psycopg2.connect(
            host=os.environ["POSTGRES_HOST"],
            user=os.environ["POSTGRES_USER"],
            password=os.environ["POSTGRES_PASSWORD"],
            dbname=os.environ["POSTGRES_DATABASE"]
        )
        cur = conn.cursor()
        cur.execute("""
            UPDATE file_tasks SET status=%s, updated_at=NOW() WHERE id=%s;
        """, (status, task_id))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Error updating file_tasks status: {e}")

def _estimate_content_size_bytes(fcontent) -> int | None:
    """
    Best-effort byte-size estimation for incoming file content.
    """
    try:
        if isinstance(fcontent, (bytes, bytearray)):
            return len(fcontent)
        if isinstance(fcontent, str):
            # try base64 first
            try:
                return len(base64.b64decode(fcontent, validate=True))
            except Exception:
                # fallback to utf-8 encoding size
                return len(fcontent.encode("utf-8"))
        # last resort
        return len(bytes(fcontent))
    except Exception:
        return None   
       

def _format_size_with_unit(size_bytes: int | None) -> str | None:
    """
    Convert a byte size into a human-readable string with units.
    Returns values with up to two decimal places using binary units (KB, MB, GB).

    Examples:
    - 1536 -> "1.50 KB"
    - 1048576 -> "1.00 MB"
    - 1073741824 -> "1.00 GB"
    """
    try:
        if size_bytes is None:
            return None
        # Use binary measurement (KiB, MiB, GiB) but display as KB/MB/GB
        KB = 1024
        MB = KB * 1024
        GB = MB * 1024

        if size_bytes < KB:
            # Keep bytes as-is; although request focuses on KB/MB/GB,
            # small files will still be accurately represented.
            return f"{size_bytes} Bytes"
        elif size_bytes < MB:
            return f"{size_bytes / KB:.2f} KB"
        elif size_bytes < GB:
            return f"{size_bytes / MB:.2f} MB"
        else:
            return f"{size_bytes / GB:.2f} GB"
    except Exception:
        # Fallback to raw bytes string if something goes wrong
        return str(size_bytes) if size_bytes is not None else None
 
# Reusable upload function (keeps same functionality as your original upload_files_and_get_urls)
# async def upload_files_and_get_urls(container_name: str, file_path: str, file_names: list, file_contents: list, expiry_years: int = 10):
   
#     print("File names to upload:", file_names)
#     print("File contents types:", [type(fc) for fc in file_contents])
#     print("File contents sizes:", [len(fc) if isinstance(fc, (bytes, str)) else 'N/A' for fc in file_contents])
#     """
#     Uploads files to Azure Blob Storage and returns their long-lived SAS download URLs.
 
#     Args:
#         container_name (str): Name of the container where file will be uploaded.
#         file_path (str): Path of the uploaded file.
#         file_names (list): List of file names to upload.
#         file_contents (list): List of file contents (bytes or str).
#         expiry_years (int): Years until the SAS token expires (default: 10).
 
#     Returns:
#         dict: {file_name: download_url or error}
#     """
#     connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
 
#     if not connection_string or not container_name:
#         return {"error": "Azure Blob Storage configuration is missing."}
 
#     if not (len(file_names) == len(file_contents)):
#         return {"error": "file_names and file_contents must have the same length."}
 
#     try:
#         blob_service_client = BlobServiceClient.from_connection_string(connection_string)
#         container_client = blob_service_client.get_container_client(container_name)
#         result = {}
#         account_name = str(blob_service_client.account_name) if blob_service_client.account_name else ""
#         if not account_name:
#             return {"error": "Could not determine Azure Storage account name."}
#         for fname, fcontent in zip(file_names, file_contents):
#             # Normalize input content to raw bytes
#             original_type = type(fcontent)
 
#             # Handle data URL strings
#             if isinstance(fcontent, str):
#                 if fcontent.startswith("data:"):
#                     try:
#                         _, b64data = fcontent.split(",", 1)
#                         fcontent = b64data
#                     except ValueError:
#                         pass
#                 # Try base64 decode; if that fails, treat as UTF-8 text bytes
#                 try:
#                     fcontent = base64.b64decode(fcontent, validate=True)
#                 except Exception:
#                     fcontent = fcontent.encode("utf-8")
 
#             # Detect and decode base64 provided as bytes
#             elif isinstance(fcontent, (bytes, bytearray)):
#                 sample = bytes(fcontent[:10])
#                 is_ascii_like = all(32 <= b <= 126 or b in (9, 10, 13) for b in sample)
#                 if is_ascii_like or sample.startswith(b"JVBER"):
#                     try:
#                         decoded = base64.b64decode(fcontent, validate=True)
#                         if decoded.startswith(b"%PDF"):
#                             fcontent = decoded
#                     except Exception:
#                         pass
#             else:
#                 # Fallback conversion
#                 try:
#                     fcontent = bytes(fcontent)
#                 except Exception:
#                     return {"error": f"Unsupported content type for {fname}: {original_type}"}
 
#             # Choose content type based on file extension
#             ext = os.path.splitext(fname)[1].lower()

#             # For DOC/DOCX specifically, ensure we upload raw binary bytes (not base64 text)
#             if ext in (".doc", ".docx"):
#                 try:
#                     if isinstance(fcontent, str):
#                         if fcontent.startswith("data:"):
#                             b64_part = fcontent.split("base64,", 1)[-1]
#                             fcontent = base64.b64decode(b64_part, validate=True)
#                         else:
#                             try:
#                                 fcontent = base64.b64decode(fcontent, validate=True)
#                             except Exception:
#                                 fcontent = fcontent.encode("utf-8")
#                     elif isinstance(fcontent, (bytes, bytearray)):
#                         # If bytes look like ASCII base64, try decoding generically
#                         sample = bytes(fcontent[:12])
#                         is_ascii_like = all(32 <= b <= 126 or b in (9, 10, 13) for b in sample)
#                         if is_ascii_like:
#                             try:
#                                 decoded = base64.b64decode(fcontent, validate=False)
#                                 # Accept decoded if it matches DOCX ZIP (PK..) or DOC OLE (D0 CF 11 E0 ...)
#                                 if (
#                                     decoded.startswith(b"PK\x03\x04") or
#                                     decoded.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1")
#                                 ):
#                                     fcontent = decoded
#                             except Exception:
#                                 # fall through, keep original bytes
#                                 pass
#                         # Ensure type is bytes
#                         fcontent = bytes(fcontent)
#                     else:
#                         fcontent = bytes(fcontent)
#                 except Exception:
#                     fcontent = bytes(fcontent) if isinstance(fcontent, (bytes, bytearray)) else str(fcontent).encode("utf-8")

#             if ext == ".pdf":
#                 content_settings = ContentSettings(content_type="application/pdf")
#             elif ext in (".txt", ".log"):
#                 content_settings = ContentSettings(content_type="text/plain")
#             elif ext == ".docx":
#                 content_settings = ContentSettings(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
#             elif ext == ".doc":
#                 content_settings = ContentSettings(content_type="application/msword")
#             else:
#                 content_settings = ContentSettings(content_type="application/octet-stream")
 
#             # Upload
#             blob_path = file_path + f"/{fname}"
#             blob_client = container_client.get_blob_client(blob_path)
#             try:
#                 blob_client.upload_blob(fcontent, overwrite=True, content_settings=content_settings)
#                 expiry = datetime.now() + timedelta(days=365 * expiry_years)
#                 sas_token = generate_blob_sas(
#                     account_name=account_name,
#                     container_name=container_name,
#                     blob_name=blob_path,
#                     account_key=blob_service_client.credential.account_key,
#                     permission=BlobSasPermissions(read=True),
#                     expiry=expiry,
#                     # Force download in browser
#                     content_disposition=f'attachment; filename="{os.path.basename(blob_path)}"'
#                 )
#                 download_url = f"https://{account_name}.blob.core.windows.net/{container_name}/{blob_path}?{sas_token}"
#                 print(f"Uploaded {fname} to {download_url}")
#                 result[fname] = download_url
#                 print(f"Upload successful for {fname}")
#             except Exception as e:
#                 result[fname] = f"Error: {str(e)}, Fcontent: {type(fcontent)}"
#         return result
#     except Exception as e:
#         return {"error": str(e)}

async def upload_files_and_get_urls(
    container_name: str,
    file_path: str,
    file_names: list,
    file_contents: list,
    expiry_years: int = 10
):
    """
    Uploads files to cloud storage and returns long‑lived download URLs.
    Uses provider‑agnostic storage_config.storage_client (Azure/S3)
    without altering original functionality.
    """

    import logging
    logger = logging.getLogger("kbAdapterTool.upload_files_and_get_urls")

    # ---- Validation ----
    if not storage_config.storage_client:
        return {"error": "Storage client not initialized"}

    if not container_name:
        return {"error": "Container/bucket name is required."}

    if len(file_names) != len(file_contents):
        return {"error": "file_names and file_contents must have the same length."}

    try:
        result = {}
        expires_in_seconds = 365 * 24 * 60 * 60 * expiry_years

        for fname, fcontent in zip(file_names, file_contents):
            try:
                original_type = type(fcontent)

                # -------------------------
                # Normalize content → raw bytes
                # -------------------------
                if isinstance(fcontent, str):
                    # Handle strings that may be base64
                    if fcontent.startswith("data:"):
                        try:
                            _, b64data = fcontent.split(",", 1)
                            fcontent = b64data
                        except:
                            pass

                    try:
                        fcontent = base64.b64decode(fcontent, validate=False)
                    except Exception:
                        fcontent = fcontent.encode("utf-8")

                elif isinstance(fcontent, (bytes, bytearray)):
                    # Attempt base64 decode only if bytes look ASCII-like
                    sample = bytes(fcontent[:10])
                    is_ascii = all(32 <= b <= 126 or b in (9, 10, 13) for b in sample)
                    if is_ascii:
                        try:
                            decoded = base64.b64decode(fcontent, validate=False)
                            fcontent = decoded
                        except Exception:
                            pass
                    fcontent = bytes(fcontent)

                else:
                    # fallback → convert to bytes
                    try:
                        fcontent = bytes(fcontent)
                    except:
                        result[fname] = f"Error: Unsupported content type {original_type}"
                        continue

                # -------------------------
                # Upload file via unified adapter
                # -------------------------
                blob_key = file_path + f"/{fname}"

                storage_config.storage_client.put_bytes(
                    container=container_name,
                    key=blob_key,
                    data=fcontent,
                    overwrite=True
                )

                # -------------------------
                # Generate presigned URL (Azure or S3)
                # -------------------------
                url = storage_config.storage_client.generate_presigned_url(
                    container=container_name,
                    key=blob_key,
                    expires_in=expires_in_seconds,
                    method="GET"
                )

                result[fname] = url

            except Exception as e:
                result[fname] = f"Error: {str(e)}"

        return result

    except Exception as e:
        return {"error": str(e)}   
 
# async def lightrag_indexing_tool(
#     container_name: Optional[str] = None, 
#     domain: Optional[str] = None, 
#     kb_name: Optional[str] = None, 
#     file_path: Optional[str] = None
# ) -> dict:
#     """
#     Index all .txt files in the specified domain/KB directory.
#     """
#     try:
#         rag = await initialize_rag(domain=domain, kb_name=kb_name)
#         if not file_path:
#             return {"error": "file_path is required"}
#         ext = file_path.lower().split('.')[-1]
#         content = None
#         # Read file from Azure Blob Storage
#         connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
#         blob_service_client = BlobServiceClient.from_connection_string(connection_string)
#         print(file_path)
#         container_client = blob_service_client.get_container_client(container_name or "")
#         blob_client = container_client.get_blob_client(file_path)
#         blob_data = blob_client.download_blob().readall()
#         print(blob_data[:100])
#         if ext == 'txt':
#             content = blob_data.decode('utf-8', errors='ignore')
#         elif ext == 'pdf':
#             # Set up credentials and client
#             try:
#                 endpoint = os.getenv('AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT')
#                 api_key = os.getenv('AZURE_DOCUMENT_INTELLIGENCE_KEY')
#                 if not endpoint or not api_key:
#                     return {"error": "Azure Document Intelligence endpoint or key not set in environment variables."}
#                 doc_client = DocumentIntelligenceClient(endpoint, AzureKeyCredential(api_key))
 
#                 # Pass the stream directly to Document Intelligence
#                 poller = doc_client.begin_analyze_document(
#                     "prebuilt-read",
#                     body=AnalyzeDocumentRequest(bytes_source=blob_data),
#                     locale="en-US"
#                 )
#                 result = poller.result()
 
#                 content = result.content
#             except Exception as e:
#                 return {"error": f"Failed to process PDF with Document Intelligence: {e}"}
#         elif ext in ['docx']:
#             # Ensure blob data represents raw DOCX/DOC bytes (not base64 text)
#             def _ensure_doc_bytes(data: bytes | str) -> bytes:
#                 try:
#                     if isinstance(data, (bytes, bytearray)):
#                         sample = bytes(data[:32])
#                         is_ascii_like = all(32 <= b <= 126 or b in (9, 10, 13) for b in sample)
#                         if is_ascii_like:
#                             try:
#                                 decoded = base64.b64decode(data, validate=True)
#                                 if decoded.startswith(b"PK\x03\x04"):
#                                     return decoded
#                             except Exception:
#                                 pass
#                         return bytes(data)
#                     elif isinstance(data, str):
#                         if data.startswith("data:"):
#                             try:
#                                 b64_part = data.split("base64,", 1)[-1]
#                                 return base64.b64decode(b64_part, validate=True)
#                             except Exception:
#                                 return data.encode("utf-8")
#                         try:
#                             return base64.b64decode(data, validate=True)
#                         except Exception:
#                             return data.encode("utf-8")
#                     else:
#                         return bytes(data)
#                 except Exception:
#                     return bytes(data) if isinstance(data, (bytes, bytearray)) else str(data).encode("utf-8")

#             doc_bytes = _ensure_doc_bytes(blob_data)
#             doc = Document(io.BytesIO(doc_bytes))
#             content = "\n".join([p.text for p in doc.paragraphs])
#         elif ext == 'doc':
#             # Attempt Windows COM-based conversion (.doc -> .docx) if Word is available
#             def _convert_doc_to_docx_bytes(doc_bytes: bytes) -> bytes | None:
#                 try:
#                     import win32com.client as win32
#                     import pythoncom
#                     with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as f_in:
#                         f_in.write(doc_bytes)
#                         in_path = f_in.name
#                     out_path = in_path[:-4] + '.docx'
#                     try:
#                         pythoncom.CoInitialize()
#                         word = win32.Dispatch('Word.Application')
#                         word.Visible = False
#                         doc = word.Documents.Open(in_path)
#                         wdFormatXMLDocument = 12
#                         doc.SaveAs(out_path, FileFormat=wdFormatXMLDocument)
#                         doc.Close(False)
#                         word.Quit()
#                         pythoncom.CoUninitialize()
#                         with open(out_path, 'rb') as f_out:
#                             return f_out.read()
#                     finally:
#                         try:
#                             import os
#                             if os.path.exists(in_path):
#                                 os.remove(in_path)
#                             if os.path.exists(out_path):
#                                 os.remove(out_path)
#                         except Exception:
#                             pass
#                 except Exception:
#                     return None

#             hdr = bytes(blob_data[:8]) if isinstance(blob_data, (bytes, bytearray)) else b''
#             if hdr.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1") or hdr.startswith(b"0M8R4KG"):
#                 converted = _convert_doc_to_docx_bytes(blob_data if isinstance(blob_data, (bytes, bytearray)) else blob_data.encode('utf-8'))
#                 if converted:
#                     doc = Document(io.BytesIO(converted))
#                     content = "\n".join([p.text for p in doc.paragraphs])
#                 else:
#                     # Fallback: try extracting readable ASCII sequences, so indexing doesn’t stop
#                     try:
#                         import re
#                         raw = blob_data if isinstance(blob_data, (bytes, bytearray)) else str(blob_data).encode('utf-8', errors='ignore')
#                         text = raw.decode('latin-1', errors='ignore')
#                         blocks = re.findall(r"[\x20-\x7E]{5,}", text)
#                         content = "\n".join(blocks)
#                         if not content or len(content.strip()) < 20:
#                             return {"error": "Legacy .doc detected and conversion unavailable. Could not extract sufficient text. Please convert to .docx and re-upload."}
#                     except Exception:
#                         return {"error": "Legacy .doc detected and conversion unavailable. Could not extract text reliably. Please convert to .docx and re-upload."}
#             else:
#                 return {"error": "Unknown .doc content format. Only .docx is supported for indexing."}
#         else:
#             return {"error": f"Unsupported file type: {ext}"}
#         def chunk_text(text, chunk_size=2000):
#             # Simple chunking by character count
#             return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
 
#         chunks = chunk_text(content, chunk_size=2000)
#         for idx, chunk in enumerate(chunks):
#             await rag.ainsert(input=chunk, file_paths=[file_path])
#             #await ctx.debug(f"Progress: {idx+1}/{len(chunks)}")
#         return {"status": "success", "file": file_path, "chunks": len(chunks)}
#     except Exception as e:
#         return {"error": str(e)}
    
# async def lightrag_indexing_tool(
#     container_name: Optional[str] = None, 
#     domain: Optional[str] = None, 
#     kb_name: Optional[str] = None, 
#     file_path: Optional[str] = None,
#     provider: Optional[str] = os.getenv("DOC_EXTRACT_PROVIDER"),  # adapter-based
#     locale: Optional[str] = "en-US",
# ) -> dict:
#     """
#     Index a single file using adapter-based storage + document extraction.
#     Functionality unchanged; only adapters applied.
#     """

#     try:
#         if not storage_config.storage_client:
#             return {"error": "Storage client not initialized"}

#         if not file_path:
#             return {"error": "file_path is required"}

#         # -------------------------
#         # Initialize RAG (unchanged)
#         # -------------------------
#         rag = await initialize_rag(domain=domain, kb_name=kb_name)

#         # -------------------------
#         # Fetch file bytes (ADAPTER)
#         # -------------------------
#         _container = container_name or storage_config.container_name
#         if not _container:
#             return {"error": "Container name not configured"}

#         blob_data = storage_config.storage_client.get_bytes(
#             container=_container,
#             key=file_path
#         )

#         ext = file_path.lower().split('.')[-1]
#         content = None

#         # -------------------------
#         # TXT (unchanged)
#         # -------------------------
#         if ext == "txt":
#             content = blob_data.decode("utf-8", errors="ignore")

#         # -------------------------
#         # DOCX / DOC (unchanged logic)
#         # -------------------------
#         elif ext in ("docx", "doc"):
#             doc = Document(io.BytesIO(blob_data))
#             content = "\n".join(p.text for p in doc.paragraphs)

#         # -------------------------
#         # PDF / Images → DocReader (ADAPTER)
#         # -------------------------
#         else:
#             reader = DocReader(provider=provider)
#             content = await reader.read(
#                 bytes=blob_data,
#                 ext=ext,
#                 locale=locale,
#                 filename=file_path
#             )

#         if not content or not content.strip():
#             return {"error": "No content extracted from file"}

#         # -------------------------
#         # Chunking (unchanged)
#         # -------------------------
#         def chunk_text(text, chunk_size=2000):
#             return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

#         chunks = chunk_text(content, chunk_size=2000)

#         # -------------------------
#         # Index chunks (unchanged)
#         # -------------------------
#         for idx, chunk in enumerate(chunks):
#             await rag.ainsert(input=chunk, file_paths=[file_path])

#         return {
#             "status": "success",
#             "file": file_path,
#             "chunks": len(chunks),
#             "provider": provider,
#         }

#     except Exception as e:
#         return {"error": str(e)}

async def lightrag_indexing_tool(
    container_name: Optional[str] = None,
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    file_path: Optional[str] = None,
    provider: Optional[str] = os.getenv("DOC_EXTRACT_PROVIDER"),
    locale: Optional[str] = "en-US",
) -> dict:
    """
    Adapter-based file indexing with deep logging.
    """
    print("\n===============================")
    print("🔵 lightrag_indexing_tool CALLED")
    print("===============================")
    print("container:", container_name)
    print("domain:", domain, "kb_name:", kb_name)
    print("file_path:", file_path)
    print("provider:", provider)
    print("===============================\n")

    try:
        # Storage config check
        if not storage_config.storage_client:
            print("❌ storage_client NOT initialized!")
            return {"error": "Storage client not initialized"}

        if not file_path:
            print("❌ file_path missing")
            return {"error": "file_path is required"}

        # -------------------------
        # INIT RAG
        # -------------------------
        print("🔧 Initializing RAG...")
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        print("✅ RAG initialized.")

        # -------------------------
        # FETCH FILE BYTES
        # -------------------------
        print("📥 Fetching file bytes from storage adapter...")
        _container = container_name or storage_config.container_name
        print("→ Effective container:", _container)

        blob_data = storage_config.storage_client.get_bytes(
            container=_container,
            key=file_path
        )

        if blob_data is None:
            print("❌ storage_client.get_bytes returned None")
            return {"error": "Failed to fetch file bytes"}

        print(f"✅ File fetched: {len(blob_data)} bytes")
        ext = file_path.lower().split('.')[-1]
        print("📄 Detected extension:", ext)

        # -------------------------
        # FILE TYPE HANDLING
        # -------------------------
        content = None

        if ext == "txt":
            print("📄 TXT file detected → decoding...")
            content = blob_data.decode("utf-8", errors="ignore")
            print("TXT decode successful.")

        elif ext in ("docx", "doc"):
            print("📄 DOCX/DOC file detected → using python-docx...")
            try:
                doc = Document(io.BytesIO(blob_data))
                content = "\n".join(p.text for p in doc.paragraphs)
                print("DOCX/DOC extraction successful. Character count:", len(content))
            except Exception as e:
                print("❌ DOCX/DOC extraction failed:", e)
                return {"error": f"Failed to process DOCX/DOC: {e}"}

        else:
            # Adapter-based DocReader for pdf / images / others
            print(f"📄 Using DocReader provider={provider} for '{ext}'...")
            reader = DocReader(provider=provider)

            try:
                content = await reader.read(
                    bytes=blob_data,
                    ext=ext,
                    locale=locale,
                    filename=file_path
                )
                print("DocReader extraction completed. Character count:", 
                      len(content) if content else "NULL")
            except Exception as e:
                tb = traceback.format_exc()
                print("❌ DocReader failed:", e)
                print(tb)
                return {"error": f"DocReader failed: {e}"}

        if not content or not content.strip():
            print("❌ Extracted content empty or whitespace")
            return {"error": "No content extracted from file"}

        # -------------------------
        # CHUNKING
        # -------------------------
        print("🔪 Starting chunking...")

        def chunk_text(text, chunk_size=2000):
            return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

        chunks = chunk_text(content, 2000)

        if not chunks:
            print("❌ No chunks produced")
            return {"error": "Could not chunk document"}

        print(f"✅ Chunking complete: {len(chunks)} chunks")

        # -------------------------
        # INDEXING
        # -------------------------
        print("📘 Starting insertion into RAG...")

        for idx, chunk in enumerate(chunks):
            print(f"➡️ Inserting chunk {idx+1}/{len(chunks)} ({len(chunk)} chars)")
            try:
                await rag.ainsert(input=chunk, file_paths=[file_path])
                print(f"   ✔ Chunk {idx+1} inserted")
            except Exception as insert_err:
                tb = traceback.format_exc()
                print(f"❌ ERROR inserting chunk {idx+1}: {insert_err}")
                print(tb)
                return {"error": f"Insert failed at chunk {idx+1}: {insert_err}"}

        print("🎉 ALL CHUNKS INSERTED SUCCESSFULLY")

        return {
            "status": "success",
            "file": file_path,
            "chunks": len(chunks),
            "provider": provider,
        }

    except Exception as e:
        tb = traceback.format_exc()
        print("❌ UNHANDLED EXCEPTION in lightrag_indexing_tool:", e)
        print(tb)
        return {"error": str(e)}

# Combined MCP tool: upload + indexing (both background)
# @mcp.tool()
# async def upload_and_index_tool(
#     ctx: Context,
#     workspace_id: Optional[str] = None,
#     container_name: Optional[str] = None,
#     upload_path: Optional[str] = None,
#     file_names: Optional[List[str]] = None,
#     file_contents: Optional[List[bytes]] = None,
#     domain: Optional[str] = None,
#     kb_name: Optional[str] = None,
#     user_id: Optional[str] = None,
#     expiry_years: int = 10
# ) -> dict:
#     """
#     Orchestrates uploading files to Azure Blob Storage and then indexing them.
#     - Creates a single task row (file_tasks) with status 'uploading'
#     - Immediately returns to client with task_id (avoids MCP timeout)
#     - Runs the upload in background. On upload success => status 'uploaded'
#     - Then starts indexing in background (status 'indexing') using the same indexing logic you had
#     - Updates status to 'indexed' or 'failed'
#     """
#     if user_id is None:
#         return {"status": "error", "error": "user_id cannot be null"}

#     # --- JWT-based authentication and workspace-user mapping check (copied from list_workspace_users) ---
#     # Validate workspace_id presence
#     # if not workspace_id:
#     #     await ctx.debug("workspace_id is required for authentication.")
#     #     return {"error": "workspace_id is required for authentication."}

#     # # Validate user access to workspace
#     # valid, err = validate_user_workspace_access(workspace_id=workspace_id)
#     # if not valid:
#     #     await ctx.debug(f"User not mapped to workspace: {err}")
#     #     return {"error": err}

#     # # Enforce JWT-based access: only allow if user is mapped to the workspace
#     # request = request_var.get(None)
#     # if not request or not hasattr(request.state, "jwt_claims"):
#     #     await ctx.debug("Unauthorized: JWT claims not found in request context")
#     #     return {"error": "Unauthorized: JWT claims not found in request context"}
#     # claims = request.state.jwt_claims
#     # jwt_user_id = claims.get("user_id") or claims.get("sub")
#     # if not jwt_user_id:
#     #     await ctx.debug("Unauthorized: user_id not found in token claims")
#     #     return {"error": "Unauthorized: user_id not found in token claims"}

#     jwt_user_id = user_id

#     # Check if user is mapped to this workspace
#     session = Session()
#     try:
#         user_map = session.query(UserMap).filter_by(workspace_id=workspace_id, user_id=jwt_user_id, is_active=True).first()
#         if not user_map:
#             session.close()
#             await ctx.debug("You are not authorized to access this workspace.")
#             return {"error": "You are not authorized to access this workspace."}
#     except Exception as e:
#         session.close()
#         await ctx.debug(f"Error during workspace-user mapping check: {e}")
#         return {"error": str(e)}
#     finally:
#         pass

#     # Validate input lists for multi-file flow
#     if not file_names or not file_contents or len(file_names) != len(file_contents):
#         await ctx.debug("file_names and file_contents are required and must match in length.")
#         return {
#             "message": "Invalid input: file_names and file_contents must be provided with equal lengths.",
#             "status": "error"
#         }

#     # Create one background job per file, each with its own task row following the same state machine
#     tasks_summary = []

#     async def background_upload_then_index_single(fname: str, fcontent: bytes, fpath: str, tid: Optional[int]):
#         try:
#             if not tid:
#                 print(f"No task id created for file {fname}; aborting background flow.")
#                 return

#             # ---- BACKGROUND UPLOAD (per-file) ----
#             update_file_task_status(tid, "uploading")
#             upload_result = await upload_files_and_get_urls(
#                 container_name,
#                 upload_path or "",
#                 [fname],
#                 [fcontent],
#                 expiry_years=expiry_years,
#             )

#             # Inspect upload result for this specific file
#             per_file_error = False
#             if isinstance(upload_result, dict):
#                 if upload_result.get("error"):
#                     per_file_error = True
#                 else:
#                     v = upload_result.get(fname)
#                     if isinstance(v, str) and v.startswith("Error:"):
#                         per_file_error = True

#             if per_file_error:
#                 update_file_task_status(tid, "failed")
#                 print("Upload failed for task_id:", tid, "result:", upload_result)
#                 return

#             # Upload success
#             update_file_task_status(tid, "uploaded")
#             print("Upload complete for task_id:", tid)

#             # ---- START INDEXING (per-file) ----
#             update_file_task_status(tid, "indexing")

#             try:
#                 if not fpath:
#                     update_file_task_status(tid, "failed")
#                     return

#                 result = await lightrag_indexing_tool(
#                     container_name=container_name,
#                     domain=domain,
#                     kb_name=kb_name,
#                     file_path=fpath,
#                 )

#                 # Handle responses consistently: treat certain benign errors as indexed
#                 if isinstance(result, dict) and result.get("error"):
#                     error_msg = str(result.get("error"))
#                     if (
#                         'already exists' in error_msg
#                         or 'No new unique documents' in error_msg
#                         or 'No documents to process' in error_msg
#                     ):
#                         print("File indexing already completed or nothing to index for task_id:", tid)
#                         update_file_task_status(tid, "indexed")
#                     else:
#                         print("Indexing failed for task_id:", tid, "error:", error_msg)
#                         update_file_task_status(tid, "failed")
#                         return
#                 else:
#                     update_file_task_status(tid, "indexed")
#                     print("Indexing complete for task_id:", tid)

#             except Exception as e:
#                 # handle exceptions from indexing
#                 tb = traceback.format_exc()
#                 error_msg = str(e) or "Unknown error (exception has no message)"
#                 if 'already exists' in error_msg or 'No new unique documents' in error_msg or 'No documents to process' in error_msg:
#                     print("File indexing already completed or nothing to index for task_id:", tid)
#                     update_file_task_status(tid, "indexed")
#                 else:
#                     print(f"Error during background indexing: {error_msg}\nTraceback:\n{tb}")
#                     update_file_task_status(tid, "failed")
#                     return

#         except Exception as e:
#             # Any top-level unexpected exception in the background flow
#             tb = traceback.format_exc()
#             print(f"Unexpected background error for task_id {tid}: {e}\n{tb}")
#             if tid:
#                 update_file_task_status(tid, "failed")

#     # Kick off one background coroutine per file
#     print("Starting background upload and indexing tasks for files:", file_names)
#     print("length file_contents", len(file_contents))
#     for fname, fcontent in zip(file_names, file_contents):
#         per_file_path = f"{upload_path}/{fname}" if upload_path and fname else None
#         # Compute human-readable size with units for storage in file_tasks.file_size
#         _bytes = _estimate_content_size_bytes(fcontent)
#         estimated_size = _format_size_with_unit(_bytes)
#         tid = create_file_task_record(
#             container_name,
#             upload_path,
#             domain,
#             kb_name,
#             per_file_path,
#             workspace_id,
#             status="uploading",
#             file_size=estimated_size,
#             uploaded_by=user_id,  # use user_id as "uploaded_by"
#         )
#         tasks_summary.append({
#             "file_name": fname,
#             "file_path": per_file_path,
#             "task_id": tid,
#         })
#         asyncio.create_task(background_upload_then_index_single(fname, fcontent, per_file_path, tid))

#     # Immediately inform client that tasks started (avoid MCP timeout)
#     await ctx.debug("Upload(s) started in background. Use the returned task_ids to poll status.")
#     return {
#         "message": "Upload and indexing started in background for all files.",
#         "status": "background",
#         "tasks": tasks_summary,
#     }

@mcp.tool()
async def upload_and_index_tool(
    ctx: Context,
    workspace_id: Optional[str] = None,
    container_name: Optional[str] = None,
    upload_path: Optional[str] = None,
    file_names: Optional[List[str]] = None,
    file_contents: Optional[List[bytes]] = None,
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    user_id: Optional[str] = None,
    expiry_years: int = 10
) -> dict:
    """
    Orchestrates uploading files and indexing them with detailed logging.
    """

    print("\n====================")
    print("🟦 upload_and_index_tool CALLED")
    print("====================")
    print("workspace_id:", workspace_id)
    print("container_name:", container_name)
    print("upload_path:", upload_path)
    print("file_names:", file_names)
    print("domain:", domain, "kb_name:", kb_name)
    print("====================\n")

    if user_id is None:
        print("❌ user_id is missing")
        return {"status": "error", "error": "user_id cannot be null"}

    # ---------------------------
    # Workspace Auth
    # ---------------------------
    print("🔍 Validating workspace-user mapping...")
    session = Session()
    try:
        user_map = session.query(UserMap).filter_by(
            workspace_id=workspace_id, user_id=user_id, is_active=True
        ).first()

        if not user_map:
            session.close()
            await ctx.debug("❌ You are not authorized to access this workspace.")
            return {"error": "You are not authorized to access this workspace."}

        print("✅ Workspace-user mapping verified.")

    except Exception as e:
        session.close()
        err = f"Workspace-user mapping error: {e}"
        print("❌", err)
        await ctx.debug(err)
        return {"error": str(e)}

    # ---------------------------
    # Validate file inputs
    # ---------------------------
    print("🔍 Validating file_names & file_contents...")
    if not file_names or not file_contents or len(file_names) != len(file_contents):
        print("❌ Invalid input lists")
        await ctx.debug("file_names and file_contents length mismatch.")
        return {
            "message": "Invalid input: file_names and file_contents must match in length.",
            "status": "error"
        }
    print("✅ file_names and file_contents validated.")

    # ---------------------------
    # Summary for client
    # ---------------------------
    tasks_summary = []

    print("\n==============================")
    print("📌 Creating background tasks…")
    print("==============================\n")

    # ---------------------------
    # INTERNAL BACKGROUND FUNCTION
    # ---------------------------
    async def background_upload_then_index_single(fname: str, fcontent: bytes, fpath: str, tid: Optional[int]):
        print(f"\n---- BACKGROUND TASK STARTED for {fname} (task_id={tid}) ----")

        if not tid:
            print("❌ No task_id created. ABORTING.")
            return

        try:
            # ---------------------
            # UPLOAD START
            # ---------------------
            print(f"⬆️ Uploading file: {fname}, path: {fpath}")
            update_file_task_status(tid, "uploading")

            upload_result = await upload_files_and_get_urls(
                container_name,
                upload_path or "",
                [fname],
                [fcontent],
                expiry_years=expiry_years,
            )

            print("UPLOAD RESULT:", upload_result)

            # Check upload errors
            per_file_error = False
            if isinstance(upload_result, dict):
                if upload_result.get("error"):
                    per_file_error = True
                else:
                    v = upload_result.get(fname)
                    if isinstance(v, str) and v.startswith("Error:"):
                        per_file_error = True

            if per_file_error:
                print(f"❌ Upload failed for {fname} (task_id={tid})")
                update_file_task_status(tid, "failed")
                return

            print(f"✅ Upload succeeded for {fname} (task_id={tid})")
            update_file_task_status(tid, "uploaded")

            # ---------------------
            # INDEXING START
            # ---------------------
            print(f"📘 Starting indexing for {fname} → file_path = {fpath}")
            update_file_task_status(tid, "indexing")

            if not fpath:
                print(f"❌ fpath is None for task_id={tid}. Cannot index.")
                update_file_task_status(tid, "failed")
                return

            result = await lightrag_indexing_tool(
                container_name=container_name,
                domain=domain,
                kb_name=kb_name,
                file_path=fpath,
            )

            print("INDEXING RESULT:", result)

            # Standardize result
            if isinstance(result, dict) and result.get("error"):
                error_msg = str(result["error"])
                if (
                    "already exists" in error_msg
                    or "No new unique documents" in error_msg
                    or "No documents to process" in error_msg
                ):
                    print(f"⚠️ Indexing says file already indexed → marking success. (task_id={tid})")
                    update_file_task_status(tid, "indexed")
                else:
                    print(f"❌ Indexing FAILED for task_id={tid}: {error_msg}")
                    update_file_task_status(tid, "failed")
                return
            else:
                print(f"✅ Indexing SUCCESS for task_id={tid}")
                update_file_task_status(tid, "indexed")

        except Exception as e:
            tb = traceback.format_exc()
            print(f"❌ UNEXPECTED INDEXING ERROR for task_id={tid}: {e}\nTRACEBACK:\n{tb}")
            update_file_task_status(tid, "failed")

    # ---------------------------
    # PER-FILE CREATION OF TASKS
    # ---------------------------
    print("📦 Preparing per-file tasks...")
    for fname, fcontent in zip(file_names, file_contents):
        per_file_path = f"{upload_path}/{fname}" if upload_path and fname else None
        print(f"📝 Preparing file: {fname}, path={per_file_path}")

        # compute size
        raw_size = _estimate_content_size_bytes(fcontent)
        est_size = _format_size_with_unit(raw_size)

        print(f"   ↳ Estimated size: {est_size}")

        # create DB record
        tid = create_file_task_record(
            container_name,
            upload_path,
            domain,
            kb_name,
            per_file_path,
            workspace_id,
            status="uploading",
            file_size=est_size,
            uploaded_by=user_id,
        )

        print(f"   ↳ Task ID created: {tid}")

        tasks_summary.append({
            "file_name": fname,
            "file_path": per_file_path,
            "task_id": tid,
        })

        # start background job
        print(f"   ↳ Starting background async job for {fname}")
        asyncio.create_task(
            background_upload_then_index_single(fname, fcontent, per_file_path, tid)
        )

    print("\n==============================")
    print("🚀 All background jobs dispatched.")
    print("==============================\n")

    await ctx.debug("Upload(s) started in background. Use returned task_ids to poll status.")

    return {
        "message": "Upload and indexing started in background for all files.",
        "status": "background",
        "tasks": tasks_summary,
    }
 
 
# MCP tool to check indexing status
@mcp.tool()
async def check_specific_indexing_status(
    ctx: Context,
    task_ids: Optional[list] = None
) -> dict:
    """
    Check the status of one or more indexing tasks by task_ids (list of int or str).
    Returns a list of status dicts for each task_id.
    """
    if not task_ids:
        return {"error": "task_ids is required (list of task ids)"}
    # Allow backward compatibility: if a single int is passed, treat as list
    if isinstance(task_ids, (int, str)):
        task_ids = [task_ids]
    try:
        conn = psycopg2.connect(
            host=os.environ["POSTGRES_HOST"],
            user=os.environ["POSTGRES_USER"],
            password=os.environ["POSTGRES_PASSWORD"],
            dbname=os.environ["POSTGRES_DATABASE"]
        )
        cur = conn.cursor()
        # Prepare query for multiple ids
        format_strings = ','.join(['%s'] * len(task_ids))
        cur.execute(f"""
            SELECT id, status, created_at, updated_at, file_path, domain, kb_name, workspace_id, file_size, uploaded_by FROM file_tasks
            WHERE id IN ({format_strings})
        """, tuple(task_ids))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        found = {row[0]: row for row in rows}
        results = []
        for tid in task_ids:
            row = found.get(int(tid)) if isinstance(tid, (int, str)) and str(tid).isdigit() else None
            if row:
                task_id, status, created_at, updated_at, file_path, domain, kb_name, workspace_id, file_size, uploaded_by = row
                results.append({
                    "task_id": task_id,
                    "status": status,
                    "created_at": str(created_at),
                    "updated_at": str(updated_at),
                    "file": file_path,
                    "domain": domain,
                    "kb_name": kb_name,
                    "workspace_id": workspace_id,
                    "file_size": file_size,
                    "uploaded_by": uploaded_by,
                })
            else:
                results.append({
                    "message": "No indexing task found for the specified task_id.",
                    "task_id": tid
                })
        return {"results": results}
    except Exception as e:
        return {"error": str(e)}
    

@mcp.tool()
async def check_indexing_status_by_workspace(
    ctx: Context,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    limit: int = 100,
    include_counts: bool = True
) -> dict:
    """
    List indexing tasks for a given workspace_id.
    Returns a list of tasks (most recent first) and optional status counts.
    """
    if user_id is None:
        return {"status": "error", "error": "user_id cannot be null"}

    # --- JWT-based authentication and workspace-user mapping check (copied from list_workspace_users) ---
    if not workspace_id:
        return {"error": "workspace_id is required"}

    # Validate user access to workspace
    valid, err = validate_user_workspace_access(workspace_id=workspace_id)
    if not valid:
        await ctx.debug(f"User not mapped to workspace: {err}")
        return {"error": err}

    # Enforce JWT-based access: only allow if user is mapped to the workspace
    request = request_var.get(None)
    if not request or not hasattr(request.state, "jwt_claims"):
        await ctx.debug("Unauthorized: JWT claims not found in request context")
        return {"error": "Unauthorized: JWT claims not found in request context"}
    claims = request.state.jwt_claims
    jwt_user_id = claims.get("user_id") or claims.get("sub")
    if not jwt_user_id:
        await ctx.debug("Unauthorized: user_id not found in token claims")
        return {"error": "Unauthorized: user_id not found in token claims"}

    # Check if user is mapped to this workspace
    session = Session()
    try:
        user_map = session.query(UserMap).filter_by(workspace_id=workspace_id, user_id=jwt_user_id, is_active=True).first()
        if not user_map:
            session.close()
            await ctx.debug("You are not authorized to access this workspace.")
            return {"error": "You are not authorized to access this workspace."}
    except Exception as e:
        session.close()
        await ctx.debug(f"Error during workspace-user mapping check: {e}")
        return {"error": str(e)}
    finally:
        pass

    try:
        conn = psycopg2.connect(
            host=os.environ["POSTGRES_HOST"],
            user=os.environ["POSTGRES_USER"],
            password=os.environ["POSTGRES_PASSWORD"],
            dbname=os.environ["POSTGRES_DATABASE"]
        )
        cur = conn.cursor()

        # Optional filter by uploaded_by derived from provided user_id
        filter_uploaded_by = None
        if user_id:
            # Resolve numeric user_id to the stored uploaded_by (full name) if possible
            filter_uploaded_by = uploaded_by_username(str(user_id))

        # Build query with optional filter
        base_query = """
            SELECT id, status, created_at, updated_at, file_path, domain, kb_name, file_size, uploaded_by, container_name
            FROM file_tasks
            WHERE workspace_id = %s
        """
        params = [workspace_id]
        if filter_uploaded_by is not None:
            base_query += " AND uploaded_by = %s"
            params.append(filter_uploaded_by)

        base_query += " ORDER BY created_at DESC LIMIT %s"
        params.append(limit)

        # Fetch recent tasks for this workspace (optionally filtered by user)
        cur.execute(base_query, tuple(params))
        rows = cur.fetchall()

        # Best-effort: prepare Azure SAS prerequisites
        connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
        account_name = None
        account_key = None
        endpoint_suffix = 'core.windows.net'
        if connection_string:
            try:
                _bsc = BlobServiceClient.from_connection_string(connection_string)
                account_name = str(_bsc.account_name) if _bsc.account_name else None
            except Exception:
                account_name = None

            def _conn_val(cs: str, key: str):
                try:
                    parts = [p for p in cs.split(';') if p]
                    for p in parts:
                        if p.startswith(key + '='):
                            return p.split('=', 1)[1]
                except Exception:
                    return None
                return None

            if not account_name:
                account_name = _conn_val(connection_string, 'AccountName')
            endpoint_suffix = _conn_val(connection_string, 'EndpointSuffix') or endpoint_suffix
            account_key = _conn_val(connection_string, 'AccountKey')

        tasks = []
        for r in rows:
            # Unpack matching 10 selected columns (includes container_name)
            task_id, status, created_at, updated_at, file_path, domain, kb_name, file_size, uploaded_by, container_name = r

            # Generate per-file SAS URL if possible (non-fatal on failure)
            download_url = None
            try:
                if connection_string and account_name and account_key and container_name and file_path:
                    sas_token = generate_blob_sas(
                        account_name=account_name,
                        container_name=container_name,
                        blob_name=file_path,
                        account_key=account_key,
                        permission=BlobSasPermissions(read=True),
                        expiry=datetime.utcnow() + timedelta(days=7),
                        content_disposition=f'attachment; filename="{os.path.basename(file_path)}"'
                    )
                    download_url = f"https://{account_name}.blob.{endpoint_suffix}/{container_name}/{file_path}?{sas_token}"
            except Exception:
                download_url = None

            tasks.append({
                "task_id": task_id,
                "status": status,
                "created_at": str(created_at),
                "updated_at": str(updated_at),
                "file": file_path,
                "file_name": os.path.basename(file_path) if file_path else None,  # <- added
                "domain": domain,
                "kb_name": kb_name,
                "file_size": file_size,
                "uploaded_by": uploaded_by,  # may be full name if resolved at insert time
                "container_name": container_name,
                "download_url": download_url,
            })

        result = {
            "workspace_id": workspace_id,
            "total": len(tasks),
            "tasks": tasks
        }
        if user_id:
            result["filtered_by_user"] = filter_uploaded_by

        if include_counts:
            counts_query = """
                SELECT status, COUNT(*)
                FROM file_tasks
                WHERE workspace_id = %s
            """
            count_params = [workspace_id]
            if filter_uploaded_by is not None:
                counts_query += " AND uploaded_by = %s"
                count_params.append(filter_uploaded_by)
            counts_query += " GROUP BY status"
            cur.execute(counts_query, tuple(count_params))
            counts = {status: count for status, count in cur.fetchall()}
            result["counts"] = counts

        cur.close()
        conn.close()

        if not tasks:
            result["message"] = "No indexing tasks found for the specified workspace_id."

        return result

    except Exception as e:
        return {"error": str(e)}


@mcp.tool()
async def generate_download_urls_by_workspace(
    ctx: Context,
    workspace_id: Optional[str] = None,
    user_id: Optional[str] = None,
    limit: int = 1000,
    expiry_days: int = 7
) -> dict:
    """
    Generate SAS download URLs for all files in a workspace.
    - Pulls records from file_tasks for the given workspace_id
    - Optionally filters by uploaded_by derived from user_id
    - Returns a list of file metadata + SAS URLs
    Also returns a single ZIP SAS URL that bundles all existing files found.
    """
    if user_id is None:
        return {"status": "error", "error": "user_id cannot be null"}

    # --- JWT-based authentication and workspace-user mapping check (copied from list_workspace_users) ---
    if not workspace_id:
        return {"error": "workspace_id is required"}

    # Validate user access to workspace
    valid, err = validate_user_workspace_access(workspace_id=workspace_id)
    if not valid:
        await ctx.debug(f"User not mapped to workspace: {err}")
        return {"error": err}

    # Enforce JWT-based access: only allow if user is mapped to the workspace
    request = request_var.get(None)
    if not request or not hasattr(request.state, "jwt_claims"):
        await ctx.debug("Unauthorized: JWT claims not found in request context")
        return {"error": "Unauthorized: JWT claims not found in request context"}
    claims = request.state.jwt_claims
    jwt_user_id = claims.get("user_id") or claims.get("sub")
    if not jwt_user_id:
        await ctx.debug("Unauthorized: user_id not found in token claims")
        return {"error": "Unauthorized: user_id not found in token claims"}

    # Check if user is mapped to this workspace
    session = Session()
    try:
        user_map = session.query(UserMap).filter_by(workspace_id=workspace_id, user_id=jwt_user_id, is_active=True).first()
        if not user_map:
            session.close()
            await ctx.debug("You are not authorized to access this workspace.")
            return {"error": "You are not authorized to access this workspace."}
    except Exception as e:
        session.close()
        await ctx.debug(f"Error during workspace-user mapping check: {e}")
        return {"error": str(e)}
    finally:
        pass

    try:
        # Resolve optional user filter
        filter_uploaded_by = uploaded_by_username(str(user_id)) if user_id else None

        # Query file_tasks for the files in the workspace
        conn = psycopg2.connect(
            host=os.environ["POSTGRES_HOST"],
            user=os.environ["POSTGRES_USER"],
            password=os.environ["POSTGRES_PASSWORD"],
            dbname=os.environ["POSTGRES_DATABASE"]
        )
        cur = conn.cursor()

        base_query = """
            SELECT id, container_name, file_path, status, domain, kb_name, uploaded_by, created_at
            FROM file_tasks
            WHERE workspace_id = %s
        """
        params = [workspace_id]
        if filter_uploaded_by is not None:
            base_query += " AND uploaded_by = %s"
            params.append(filter_uploaded_by)
        base_query += " ORDER BY created_at DESC LIMIT %s"
        params.append(limit)

        cur.execute(base_query, tuple(params))
        rows = cur.fetchall()
        cur.close()
        conn.close()

        if not rows:
            return {"workspace_id": workspace_id, "files": [], "message": "No files found for this workspace."}

        # Prepare Azure client
        connection_string = os.getenv('AZURE_BLOB_STORAGE_CONNECTION_STRING')
        if not connection_string:
            return {"error": "AZURE_BLOB_STORAGE_CONNECTION_STRING is not set."}

        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        account_name = str(blob_service_client.account_name) if blob_service_client.account_name else ""
        if not account_name:
            return {"error": "Could not determine Azure Storage account name."}

        expiry = datetime.now() + timedelta(days=expiry_days)

        results = []
        files_for_zip = []  # (container_name, blob_path, file_name)
        first_container_name = None

        for (task_id, container_name, blob_path, status, domain, kb_name, uploaded_by, created_at) in rows:
            if not container_name or not blob_path:
                # skip incomplete records
                continue

            if first_container_name is None:
                first_container_name = container_name

            container_client = blob_service_client.get_container_client(container_name)
            blob_client = container_client.get_blob_client(blob_path)

            # Check existence to avoid dead links
            try:
                exists = blob_client.exists()
            except Exception:
                exists = False

            if not exists:
                results.append({
                    "task_id": task_id,
                    "domain": domain,
                    "kb_name": kb_name,
                    "file_path": blob_path,
                    "file_name": os.path.basename(blob_path),
                    "status": status,
                    "uploaded_by": uploaded_by,
                    "created_at": str(created_at),
                    "exists": False,
                    "download_url": None
                })
                continue

            # Generate read-only SAS URL
            try:
                sas_token = generate_blob_sas(
                    account_name=account_name,
                    container_name=container_name,
                    blob_name=blob_path,
                    account_key=blob_service_client.credential.account_key,
                    permission=BlobSasPermissions(read=True),
                    expiry=expiry,
                    # Force download in browser
                    content_disposition=f'attachment; filename="{os.path.basename(blob_path)}"'
                )
                download_url = f"https://{account_name}.blob.core.windows.net/{container_name}/{blob_path}?{sas_token}"
            except Exception as e:
                download_url = f"Error generating SAS: {str(e)}"

            # Track for ZIP
            files_for_zip.append((container_name, blob_path, os.path.basename(blob_path)))

            results.append({
                "task_id": task_id,
                "domain": domain,
                "kb_name": kb_name,
                "file_path": blob_path,
                "file_name": os.path.basename(blob_path),
                "status": status,
                "uploaded_by": uploaded_by,
                "created_at": str(created_at),
                "exists": True,
                "download_url": download_url
            })

        # Create a single ZIP containing all existing files (if any)
        zip_info = {
            "url": None,
            "container": None,
            "blob_path": None,
            "file_count": 0
        }

        if files_for_zip and first_container_name:
            try:
                zip_container_name = first_container_name
                zip_blob_path = f"workspace_zips/{workspace_id}/download_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"

                # Ensure zip container exists
                zip_container_client = blob_service_client.get_container_client(zip_container_name)
                try:
                    zip_container_client.get_container_properties()
                except Exception:
                    zip_container_client.create_container()

                # Build zip in a spooled temp file (spools to disk beyond threshold)
                spooled = tempfile.SpooledTemporaryFile(max_size=50 * 1024 * 1024, mode="w+b")
                with zipfile.ZipFile(spooled, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
                    for (cname, bpath, fname) in files_for_zip:
                        try:
                            cclient = blob_service_client.get_container_client(cname)
                            bclient = cclient.get_blob_client(bpath)
                            data = bclient.download_blob().readall()
                            zf.writestr(fname, data)
                        except Exception as e:
                            # Skip files that fail to download
                            continue

                spooled.seek(0)
                zip_blob_client = zip_container_client.get_blob_client(zip_blob_path)
                zip_blob_client.upload_blob(spooled, overwrite=True, content_settings=ContentSettings(content_type="application/zip"))
                spooled.close()

                # Generate SAS for ZIP
                zip_sas = generate_blob_sas(
                    account_name=account_name,
                    container_name=zip_container_name,
                    blob_name=zip_blob_path,
                    account_key=blob_service_client.credential.account_key,
                    permission=BlobSasPermissions(read=True),
                    expiry=expiry,
                    content_disposition=f'attachment; filename="{os.path.basename(zip_blob_path)}"'
                )
                zip_url = f"https://{account_name}.blob.core.windows.net/{zip_container_name}/{zip_blob_path}?{zip_sas}"

                zip_info = {
                    "url": zip_url,
                    "container": zip_container_name,
                    "blob_path": zip_blob_path,
                    "file_count": len(files_for_zip)
                }
            except Exception as e:
                zip_info = {
                    "url": f"Error creating ZIP: {str(e)}",
                    "container": first_container_name,
                    "blob_path": None,
                    "file_count": len(files_for_zip)
                }

        response = {
            "workspace_id": workspace_id,
            "total": len(results),
            "expiry_days": expiry_days,
            "files": results,
            "zip": zip_info
        }
        if user_id:
            response["filtered_by_user"] = filter_uploaded_by

        return response

    except Exception as e:
        return {"error": str(e)}
   
 
# @mcp.tool()
# async def conversation_indexing_tool(
#     ctx: Context,
#     domain: Optional[str] = None,
#     kb_name: Optional[str] = None,
#     url: Optional[str] = None
#     ) -> dict:
#     """
#     Index all user queries into the specified domain/KB directory.
#     """
#     try:
#         print("Starting crawling for URL...")
#         browser_config = BrowserConfig(verbose=True)
#         run_config = CrawlerRunConfig(
#                 markdown_generator=DefaultMarkdownGenerator(
#                 options={"ignore_links": True}
#             ),
#             # Content filtering
#             word_count_threshold=50,
#             excluded_tags=['form', 'header'],
#             exclude_external_links=True,
 
#             # Content processing
#             process_iframes=True,
#             remove_overlay_elements=True,
#         )
#         print("Browser and run config set up. Starting crawler...")
 
#         async with AsyncWebCrawler(config=browser_config) as crawler:
#             result = await crawler.arun(
#                 url=url,
#                 config=run_config
#             )
 
#         print("Crawling completed. Starting indexing...")
 
#         content = result.markdown
#         print("Content: ", content[:500])
 
#         def chunk_text(text, chunk_size=2000):
#             # Simple chunking by character count
#             return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
       
#         chunks = chunk_text(content, chunk_size=2000)
#         print("First chunk: ",chunks[0])
 
#         rag = await initialize_rag(domain=domain, kb_name=kb_name)
#         for idx, chunk in enumerate(chunks):
#             await rag.ainsert(input=chunk, file_paths=[url])
#             # await ctx.debug(f"Progress: {idx+1}/{len(chunks)}")
#         return {"status": "success", "file": url, "chunks": len(chunks)}
#     except Exception as e:
#         return {"error": str(e)}
    
@mcp.tool()
async def conversation_indexing_tool(
    ctx: Context,
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    url: Optional[str] = None
    ) -> dict:
    """
    Index all user queries into the specified domain/KB directory.
    """
    try:
        print("Starting crawling for URL...")
        browser_config = BrowserConfig(verbose=True)
 
        # Do NOT drop link text on Wikipedia. Lower the word threshold.
        run_config = CrawlerRunConfig(
            markdown_generator=DefaultMarkdownGenerator(
                options={
                    "ignore_links": False  # keep anchor text
                }
            ),
            word_count_threshold=0,       # accept short sections
            excluded_tags=['form', 'header'],
            exclude_external_links=True,
            process_iframes=True,
            remove_overlay_elements=True,
        )
        print("Browser and run config set up. Starting crawler...")
 
        async with AsyncWebCrawler(config=browser_config) as crawler:
            result = await crawler.arun(url=url, config=run_config)
 
            content = (result.markdown or "").strip()
            # Fallback retry with minimal filtering if empty
            if not content:
                print("Empty markdown, retrying with minimal filtering...")
                retry_config = CrawlerRunConfig(
                    markdown_generator=DefaultMarkdownGenerator(options={"ignore_links": False}),
                    word_count_threshold=0
                )
                retry = await crawler.arun(url=url, config=retry_config)
                content = ((retry.markdown or "") or (getattr(retry, "cleaned_text", "") or "")).strip()
 
        print("Crawling completed. Starting indexing...")
        print("Content: ", content[:500])
 
        if not content:
            return {"error": "No content extracted from URL. Try a different URL or relax crawler filters."}
 
        def chunk_text(text, chunk_size=2000):
            return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
       
        chunks = chunk_text(content, chunk_size=2000)
        if not chunks:
            return {"error": "No chunks produced from content."}
        print("First chunk: ", chunks[0][:200])
 
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        for idx, chunk in enumerate(chunks):
            await rag.ainsert(input=chunk, file_paths=[url])
        return {"status": "success", "file": url, "chunks": len(chunks)}
    except Exception as e:
        return {"error": str(e)}    
 
# @mcp.tool()
# async def ingest_kb_file(domain: str = None, kb_name: str = None, file_name: str = None, file_bytes: bytes = None) -> str:
#     """
#     Save an uploaded file to the correct KB directory.
#     """
#     try:
#         if not domain or not kb_name or not file_name or not file_bytes:
#             raise ValueError("All parameters (domain, kb_name, file_name, file_bytes) are required.")
#         kb_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'data', domain, kb_name))
#         os.makedirs(kb_dir, exist_ok=True)
#         dest_path = os.path.join(kb_dir, file_name)
#         with open(dest_path, "wb") as f:
#             f.write(file_bytes)
#         return f"File {file_name} uploaded to {kb_dir}"
#     except Exception as e:
#         return f"Error: {e}"
   
# @mcp.tool()
# async def get_kb_knowledge_graph(
#     ctx: Context,
#     domain: Optional[str],
#     kb_name: Optional[str],
#     question: Optional[str],
#     role_id: Optional[int],
#     workspace_id: Optional[int],
#     history: Optional[List] = [],
#     mode: Optional[str] = "mix",
#     user_prompt: Optional[str] = "",
#     knowledge_bases: Optional[List[str]] = None
# ) -> dict:
#     """
#     Get the knowledge graph for the specified workspace (domain_kb_name) or for provided knowledge bases.
#     Returns nodes and edges as a dict for visualization.
#     """
#     if role_id != 34:
#         digit_map = {
#             '0': 'zero', '1': 'one', '2': 'two', '3': 'three', '4': 'four',
#             '5': 'five', '6': 'six', '7': 'seven', '8': 'eight', '9': 'nine'
#         }
#         result = []
#         for c in str(workspace_id):
#             if c.isalpha():
#                 result.append(c)
#             elif c.isdigit():
#                 result.append(digit_map[c])
#             # skip non-alphanumeric characters
#         workspace_id_alpha = ''.join(result)
#         kb_name = f"{kb_name}/{workspace_id_alpha}"

#     try:
#         all_nodes = []
#         all_edges = []
#         print("[DEBUG] get_kb_knowledge_graph called with:", {
#             "domain": domain,
#             "kb_name": kb_name,
#             "question": question,
#             "history": history,
#             "mode": mode,
#             "user_prompt": user_prompt,
#             "knowledge_bases": knowledge_bases
#         })
#         # If knowledge_bases are provided, query each and merge results
#         if knowledge_bases and isinstance(knowledge_bases, list) and len(knowledge_bases) > 0:
#             print(f"[DEBUG] knowledge_bases provided: {knowledge_bases}")
#             semaphore = asyncio.Semaphore(4)  # Lower concurrency to avoid event loop overload
#             async def fetch_graph_for_kb(kb):
#                 async with semaphore:
#                     try:
#                         combined_kb_name = f"{kb_name}{kb}"
#                         print(f"[DEBUG] Initializing RAG for KB: domain={domain}, kb_name={combined_kb_name}")
#                         rag = await initialize_rag(domain=domain, kb_name=combined_kb_name)
#                         query_graph_data = await rag.aquery_data(
#                             question,
#                             param=QueryParam(
#                                 mode=mode,
#                                 top_k=5,
#                                 conversation_history=history,
#                                 user_prompt=user_prompt,
#                             )
#                         )
#                         print(f"[DEBUG] query_graph_data for KB '{kb}': {query_graph_data}")
#                         all_labels = query_graph_data["data"].get("entities", [])
#                         print(f"[DEBUG] all_labels for KB '{kb}': {all_labels}")
#                         if not all_labels:
#                             return kb, [], []
#                         label_semaphore = asyncio.Semaphore(4)
#                         async def fetch_graph(label):
#                             async with label_semaphore:
#                                 try:
#                                     print(f"[DEBUG] Fetching graph for label: {label}")
#                                     graph = await rag.get_knowledge_graph(label, 2, None)
#                                     print(f"[DEBUG] Graph for label '{label}': nodes={{}} edges={{}}".format(len(graph.nodes) if graph else 'None', len(graph.edges) if graph else 'None'))
#                                     return label, graph
#                                 except Exception as neo4j_err:
#                                     print(f"[DEBUG] Exception in fetch_graph for label '{label}': {neo4j_err}")
#                                     return label, None
#                         # Run fetch_graph in small batches to avoid timeouts
#                         batch_size = 4
#                         results = []
#                         label_names = [label['entity_name'] for label in all_labels]
#                         for i in range(0, len(label_names), batch_size):
#                             batch = label_names[i:i+batch_size]
#                             batch_results = await asyncio.gather(*(fetch_graph(label) for label in batch))
#                             results.extend(batch_results)
#                         nodes = []
#                         edges = []
#                         for idx, (label, graph) in enumerate(results):
#                             if graph:
#                                 nodes.extend([dict(node) for node in graph.nodes])
#                                 edges.extend([dict(edge) for edge in graph.edges])
#                         return kb, nodes, edges
#                     except Exception as neo4j_err:
#                         print(f"[DEBUG] Exception in fetch_graph_for_kb for KB '{kb}': {neo4j_err}")
#                         return kb, [], []
#             # Run fetch_graph_for_kb in small batches as well
#             batch_size_kb = 2
#             kb_results = []
#             for i in range(0, len(knowledge_bases), batch_size_kb):
#                 batch = knowledge_bases[i:i+batch_size_kb]
#                 batch_results = await asyncio.gather(*(fetch_graph_for_kb(kb) for kb in batch))
#                 kb_results.extend(batch_results)
#             for kb, nodes, edges in kb_results:
#                 all_nodes.extend(nodes)
#                 all_edges.extend(edges)
#             print(f"[DEBUG] Final merged nodes: {len(all_nodes)}, edges: {len(all_edges)}")
#             return {
#                 "knowledge_bases": knowledge_bases,
#                 "nodes": all_nodes,
#                 "edges": all_edges,
#             }
#         else:
#             if not domain or not kb_name:
#                 print("[DEBUG] Missing domain or kb_name")
#                 raise ValueError("Parameters 'domain', and 'kb_name' are required if knowledge_bases is not provided.")
#             print(f"[DEBUG] Initializing RAG for domain={domain}, kb_name={kb_name}")
#             rag = await initialize_rag(domain=domain, kb_name=kb_name)
#             workspace = f"{domain}_{kb_name}"
#             print(f"[DEBUG] Querying graph data with question: {question}")
#             query_graph_data = await rag.aquery_data(
#                 question,
#                 param=QueryParam(
#                     mode=mode,
#                     top_k=5,
#                     conversation_history=history,
#                     user_prompt=user_prompt,
#                 )
#             )
#             print(f"[DEBUG] query_graph_data: {query_graph_data}")
#             all_labels = query_graph_data["data"].get("entities", [])
#             print(f"[DEBUG] all_labels: {all_labels}")
#             if not all_labels:
#                 print(f"[DEBUG] No entities found for workspace {workspace}")
#                 return {workspace: {"nodes": [], "edges": []}}
#             semaphore = asyncio.Semaphore(4)
#             async def fetch_graph(label):
#                 async with semaphore:
#                     try:
#                         print(f"[DEBUG] Fetching graph for label: {label}")
#                         graph = await rag.get_knowledge_graph(label, 2, None)
#                         print(f"[DEBUG] Graph for label '{label}': nodes={len(graph.nodes) if graph else 'None'}, edges={len(graph.edges) if graph else 'None'}")
#                         return label, graph
#                     except Exception as neo4j_err:
#                         print(f"[DEBUG] Exception in fetch_graph for label '{label}': {neo4j_err}")
#                         return label, None
#             batch_size = 4
#             label_names = [label['entity_name'] for label in all_labels]
#             results = []
#             for i in range(0, len(label_names), batch_size):
#                 batch = label_names[i:i+batch_size]
#                 batch_results = await asyncio.gather(*(fetch_graph(label) for label in batch))
#                 results.extend(batch_results)
#             for idx, (label, graph) in enumerate(results):
#                 if graph:
#                     all_nodes.extend([dict(node) for node in graph.nodes])
#                     all_edges.extend([dict(edge) for edge in graph.edges])
#             print(f"[DEBUG] Final nodes: {len(all_nodes)}, edges: {len(all_edges)}")
#             return {
#                 workspace: {
#                     "nodes": all_nodes,
#                     "edges": all_edges,
#                 }
#             }
#     except Exception as e:
#         print(f"[DEBUG] get_kb_knowledge_graph failed: {e}")
#         return {"error": f"get_kb_knowledge_graph failed: {str(e)}"}
 
@mcp.tool()
async def get_kb_knowledge_graph(
    ctx: Context,
    domain: Optional[str],
    kb_name: Optional[str],
    question: Optional[str],
    role_id: Optional[int],
    workspace_id: Optional[int],
    history: Optional[List] = [],
    mode: Optional[str] = "mix",
    user_prompt: Optional[str] = "",
    knowledge_bases: Optional[List[str]] = None
) -> dict:
    """
    Normalized KG fetch using LightRAG's aquery_data() + NeptuneToNeo4jConverter,
    EXACTLY matching the provided structure while keeping all your original
    role-based logic and KB merging behavior.
    """

    try:
        # ---------------------------------------------------------
        # Role-based workspace ID rewrite (preserve EXACT logic)
        # ---------------------------------------------------------
        if role_id != 34:
            digit_map = {
                '0': 'zero', '1': 'one', '2': 'two', '3': 'three', '4': 'four',
                '5': 'five', '6': 'six', '7': 'seven', '8': 'eight', '9': 'nine'
            }
            result = []
            for c in str(workspace_id):
                if c.isalpha():
                    result.append(c)
                elif c.isdigit():
                    result.append(digit_map[c])
            workspace_id_alpha = ''.join(result)
            kb_name = f"{kb_name}/{workspace_id_alpha}"

        # ---------------------------------------------------------
        # Validate required parameters
        # ---------------------------------------------------------
        if not question:
            raise ValueError("Parameter 'question' is required.")

        # Collector for merged output
        merged_nodes: List[dict] = []
        merged_edges: List[dict] = []

        print(
            "[DEBUG] get_kb_knowledge_graph normalized structure MODE C",
            {"domain": domain, "kb_name": kb_name, "question": question}
        )

        # ---------------------------------------------------------
        # Handle MULTI-KB mode
        # ---------------------------------------------------------
        if knowledge_bases and isinstance(knowledge_bases, list) and len(knowledge_bases) > 0:
            print("[DEBUG] Multi-KB mode enabled:", knowledge_bases)

            semaphore = asyncio.Semaphore(4)

            async def fetch_normalized_for_kb(kb_suffix: str):
                """
                Per-KB retrieval using:
                  rag.aquery_data() -> converter.transform(resp["data"])
                Nodes and edges *not modified* after transform().
                """
                async with semaphore:
                    try:
                        combined_kb = f"{kb_name}{kb_suffix}"
                        print(f"[DEBUG] Initializing RAG for KB: {combined_kb}")

                        rag = await initialize_rag(domain=domain, kb_name=combined_kb)

                        param = QueryParam(
                            mode="mix",
                            top_k=5,
                            conversation_history=history,
                            user_prompt=user_prompt,
                        )

                        resp = await rag.aquery_data(question, param=param)
                        print(
                            f"[DEBUG] aquery_data keys for KB={kb_suffix}:",
                            resp.keys() if isinstance(resp, dict) else type(resp)
                        )

                        converter = NeptuneToNeo4jConverter()
                        transformed = converter.transform(resp.get("data", {}))

                        nodes = transformed.get("nodes", []) or []
                        edges = transformed.get("edges", []) or []

                        print(f"[DEBUG] KB {kb_suffix}: nodes={len(nodes)}, edges={len(edges)}")

                        return kb_suffix, nodes, edges

                    except Exception as e:
                        print(f"[DEBUG] Error in KB={kb_suffix} → {e}")
                        return kb_suffix, [], []

            # Batch KB calls
            results = []
            batch_size = 2
            for i in range(0, len(knowledge_bases), batch_size):
                batch = knowledge_bases[i:i+batch_size]
                batch_results = await asyncio.gather(
                    *(fetch_normalized_for_kb(kb) for kb in batch)
                )
                results.extend(batch_results)

            # Merge results (OPTION C BEHAVIOR)
            for kb_name_item, nodes, edges in results:
                merged_nodes.extend(nodes)
                merged_edges.extend(edges)

            print(
                "[DEBUG] Multi-KB merged result:",
                {"nodes": len(merged_nodes), "edges": len(merged_edges)}
            )

            return {
                "knowledge_bases": knowledge_bases,
                "nodes": merged_nodes,
                "edges": merged_edges
            }

        # ---------------------------------------------------------
        # SINGLE-KB MODE (Option C)
        # ---------------------------------------------------------
        if not domain or not kb_name:
            raise ValueError("Parameters 'domain' and 'kb_name' are required.")

        workspace = f"{domain}_{kb_name}"
        print(f"[DEBUG] Single-KB lookup for workspace: {workspace}")

        rag = await initialize_rag(domain=domain, kb_name=kb_name)

        param = QueryParam(
                    mode="mix",
                    top_k=5,
                    conversation_history=history,
                    user_prompt=user_prompt,
                )

        resp = await rag.aquery_data(question, param=param)
        print(
            "[DEBUG] aquery_data raw response:",
            resp.keys() if isinstance(resp, dict) else type(resp)
        )

        # Strict converter-only transformation
        converter = NeptuneToNeo4jConverter()
        transformed = converter.transform(resp.get("data", {}))

        nodes = transformed.get("nodes", []) or []
        edges = transformed.get("edges", []) or []

        print(f"[DEBUG] Final normalized KG: nodes={len(nodes)}, edges={len(edges)}")

        return {
            workspace: {
                "nodes": nodes,
                "edges": edges
            }
        }

    except Exception as e:
        print(f"[DEBUG] get_kb_knowledge_graph FAILED: {e}")
        return {"error": f"get_kb_knowledge_graph failed: {str(e)}"}

@mcp.tool()
async def insert_node_to_kg(
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    entity_name: Optional[str] = None,
    entity_data: Optional[dict] = None
):
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        node_data = {
            "entity_id": entity_name,
            "entity_type": entity_data.get("entity_type", "UNKNOWN") if entity_data else "UNKNOWN",
            "description": entity_data.get("description", "") if entity_data else "",
            "source_id": entity_data.get("source_id", "manual_creation") if entity_data else "manual_creation",
            "file_path": entity_data.get("file_path", "manual_creation") if entity_data else "manual_creation",
            "created_at": int(time.time()),
        }
        return await rag.acreate_entity(
            entity_name=entity_name,
            entity_data=node_data
        )
    except Exception as e:
        return {"error": str(e)}
 
@mcp.tool()
async def insert_edge_to_kg(
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    source_entity_name: Optional[str] = None,
    target_entity_name: Optional[str] = None,
    relation_data: Optional[dict] = None
):
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        edge_data = {
            "description": relation_data.get("description", "") if relation_data else "",
            "keywords": relation_data.get("keywords", "") if relation_data else "",
            "source_id": relation_data.get("source_id", "manual_creation") if relation_data else "manual_creation",
            "weight": float(relation_data.get("weight", 1.0)) if relation_data else 1.0,
            "file_path": relation_data.get("file_path", "manual_creation") if relation_data else "manual_creation",
            "created_at": int(time.time()),
        }
        return await rag.acreate_relation(
            source_entity=source_entity_name,
            target_entity=target_entity_name,
            relation_data=edge_data
        )
    except Exception as e:
        return {"error": str(e)}
 
@mcp.tool()
async def delete_entity_from_kg(
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    entity_name: Optional[str] = None,
):
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        return await rag.adelete_by_entity(
            entity_name=entity_name,
        )
    except Exception as e:
        return {"error": str(e)}
   
@mcp.tool()
async def delete_relation_from_kg(
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    source_entity_name: Optional[str] = None,
    target_entity_name: Optional[str] = None,
):
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        return await rag.adelete_by_relation(
            source_entity=source_entity_name,
            target_entity=target_entity_name,
        )
    except Exception as e:
        return {"error": str(e)}
 
@mcp.tool()
async def edit_entity_in_kg(
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    entity_name: Optional[str] = None,
    updated_data: Optional[dict] = None,
):
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        return await rag.aedit_entity(
            entity_name=entity_name,
            updated_data=updated_data,
            allow_rename = True
        )
    except Exception as e:
        return {"error": str(e)}
 
@mcp.tool()
async def edit_relation_in_kg(
    domain: Optional[str] = None,
    kb_name: Optional[str] = None,
    source_entity_name: Optional[str] = None,
    target_entity_name: Optional[str] = None,
    updated_data: Optional[dict] = None,
):
    try:
        rag = await initialize_rag(domain=domain, kb_name=kb_name)
        return await rag.aedit_relation(
            source_entity=source_entity_name,
            target_entity=target_entity_name,
            updated_data=updated_data,
        )
    except Exception as e:
        return {"error": str(e)}
   
@mcp.tool()
async def extract_keywords_from_query(
    user_query: Optional[str] = None,
    history: Optional[list] = None,
    node_labels: Optional[list] = None
) -> dict:
    """
    Extract keywords from a user query, considering node labels as context.
    Returns only keywords related to the user prompt and the provided node labels.
    """
    try:
        if not user_query or not isinstance(user_query, str):
            return {"error": "user_query must be a non-empty string."}
        node_labels = node_labels or []
        node_labels_str = ', '.join([str(label) for label in node_labels]) if node_labels else ''
        system_prompt = (
            f"""Given the following node labels in the knowledge graph: [{node_labels_str}].
            Extract node labels from the user query that are most relevant to these node labels.
            Return only these node labels that match or are closely related to the node labels as strings in a list.\nUser Query:"""
        )
        keywords = await llm_model_func(prompt=system_prompt+user_query)
        return {"keywords": keywords}
    except Exception as e:
        return {"error": str(e)}
   
@mcp.tool()
def get_indexed_file_names(domain,kb_name):
    """
    Get the name of indexed file names
    """
    lightrag_database = ''.join(char for char in f"{domain}{kb_name}" if char.isalpha())
    with open(f'data/{lightrag_database}/kv_store_text_chunks.json', 'r', encoding='utf-8') as f:
         meta = json.load(f)
 
    doc_dict = {}
 
    for key, value in meta.items():
        doc_id = value.get('full_doc_id')
        file_path = value.get('file_path', '')
        file_name = os.path.basename(file_path)
 
        if doc_id:
            if file_name in doc_dict:
                # Append to existing list
                doc_dict[file_name].append(doc_id)
            else:
                # Create new list
                doc_dict[file_name] = [doc_id]
 
            print(f"Doc ID: {doc_id} | File Name: {file_name}")
 
    print("Doc ID to File Name Dictionary:", doc_dict)
    return doc_dict
 
@mcp.tool()
async def delete_by_doc_id(doc_id):
    """
    Delete a document by doc_id using LightRAG
    """
    try:
        rag = await initialize_rag()
        response = await rag.adelete_by_doc_id(doc_id)
        await rag.aclear_cache()
        print(f"Document with doc_id '{doc_id}' deleted successfully.")
        return response
    except Exception as e:
        print(f"Error deleting document with doc_id '{doc_id}': {e}")
 
@mcp.tool()
async def clear_cache(doc_id):
     try:
        rag = await initialize_rag()
        await rag.aclear_cache()
     except Exception as e:
        print(f"Error clearing cache': {e}")
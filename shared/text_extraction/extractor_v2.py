"""
Dedicated service for robust text extraction by file type - Multi-provider OCR support.

This is the updated version that uses the multi-provider OCR adapter system.
Shared between indexer-service and kb-rest-service.
"""

import asyncio
import io
import logging
import os
import zipfile
from typing import Optional

import PyPDF2
import pdfplumber
from docx import Document

from shared.adapters.ocr import OCRAdapter, NoOpOCRAdapter
from .decoders import normalize_file_bytes
from .models import TextExtractionError, TextExtractionResult

logger = logging.getLogger(__name__)


class BaseExtractor:
    """Base extractor contract."""

    supported_extensions: tuple[str, ...] = ()

    def can_extract(self, extension: str) -> bool:
        return extension.lower() in self.supported_extensions

    async def extract(self, file_bytes: bytes, file_path: str) -> TextExtractionResult:
        raise NotImplementedError


class PlainTextExtractor(BaseExtractor):
    supported_extensions = (".txt", ".md", ".csv", ".json", ".xml", ".yaml", ".yml", ".log")

    async def extract(self, file_bytes: bytes, file_path: str) -> TextExtractionResult:
        for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
            try:
                text = file_bytes.decode(encoding)
                return TextExtractionResult(
                    text=text,
                    extractor="plain_text",
                    metadata={"encoding": encoding},
                )
            except UnicodeDecodeError:
                continue

        raise TextExtractionError(
            "Unable to decode text file with supported encodings",
            file_path=file_path,
        )


class PdfExtractor(BaseExtractor):
    """PDF extractor with multi-provider OCR fallback."""

    supported_extensions = (".pdf",)

    def __init__(self, ocr_adapter: Optional[OCRAdapter] = None) -> None:
        self._ocr = ocr_adapter or NoOpOCRAdapter()

    async def _extract_pdfplumber(self, file_bytes: bytes) -> str:
        def _read() -> str:
            chunks = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    chunks.append(page.extract_text() or "")
            return "\n\n".join(chunks)

        return await asyncio.to_thread(_read)

    async def _extract_pypdf2(self, file_bytes: bytes) -> str:
        def _read() -> str:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            chunks = []
            for page in reader.pages:
                chunks.append(page.extract_text() or "")
            return "\n\n".join(chunks)

        return await asyncio.to_thread(_read)

    async def _extract_with_ocr(self, file_bytes: bytes, file_path: str) -> str:
        """Use OCR adapter as fallback."""
        if not self._ocr.is_configured:
            raise TextExtractionError(
                f"OCR not configured (provider: {self._ocr.provider_name}). "
                f"Cannot extract text from scanned PDF: {file_path}"
            )

        try:
            return await self._ocr.extract_text(file_bytes, file_path)
        except Exception as e:
            raise TextExtractionError(
                f"OCR extraction failed with {self._ocr.provider_name}: {str(e)}",
                file_path=file_path,
                cause=e,
            ) from e

    async def extract(self, file_bytes: bytes, file_path: str) -> TextExtractionResult:
        # Try pdfplumber first
        first_pass = await self._extract_pdfplumber(file_bytes)
        if len(first_pass.strip()) >= 100:
            return TextExtractionResult(text=first_pass, extractor="pdfplumber")

        # Try PyPDF2 second
        second_pass = await self._extract_pypdf2(file_bytes)
        if len(second_pass.strip()) >= 100:
            return TextExtractionResult(text=second_pass, extractor="pypdf2")

        # Fall back to OCR
        fallback = await self._extract_with_ocr(file_bytes, file_path)
        if len(fallback.strip()) >= 10:
            return TextExtractionResult(
                text=fallback,
                extractor=f"ocr_{self._ocr.provider_name}",
            )

        raise TextExtractionError("Could not extract meaningful text from PDF", file_path=file_path)


class DocxExtractor(BaseExtractor):
    supported_extensions = (".docx",)

    @staticmethod
    def _validate_docx_payload(file_bytes: bytes, file_path: str) -> None:
        if not file_bytes.startswith(b"PK"):
            raise TextExtractionError(
                "Invalid DOCX payload: expected zipped Office document",
                file_path=file_path,
            )

        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes), "r") as archive:
                if "word/document.xml" not in archive.namelist():
                    raise TextExtractionError(
                        "Invalid DOCX payload: missing word/document.xml",
                        file_path=file_path,
                    )
        except zipfile.BadZipFile as exc:
            raise TextExtractionError(
                "Invalid DOCX payload: corrupt zip structure",
                file_path=file_path,
                cause=exc,
            ) from exc

    async def extract(self, file_bytes: bytes, file_path: str) -> TextExtractionResult:
        self._validate_docx_payload(file_bytes, file_path)

        def _read() -> str:
            doc = Document(io.BytesIO(file_bytes))
            parts: list[str] = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        parts.append(" | ".join(row_cells))

            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        parts.append(paragraph.text)
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        parts.append(paragraph.text)

            return "\n\n".join(parts)

        text = await asyncio.to_thread(_read)
        if len(text.strip()) < 10:
            raise TextExtractionError("DOCX text is empty after extraction", file_path=file_path)

        return TextExtractionResult(text=text, extractor="docx")


class DocExtractor(BaseExtractor):
    """Legacy .doc extractor with multi-provider OCR fallback."""

    supported_extensions = (".doc",)

    def __init__(self, ocr_adapter: Optional[OCRAdapter] = None) -> None:
        self._ocr = ocr_adapter or NoOpOCRAdapter()

    async def extract(self, file_bytes: bytes, file_path: str) -> TextExtractionResult:
        # Legacy Word binary files should start with OLE signature bytes.
        is_binary_doc = file_bytes.startswith(bytes.fromhex("D0CF11E0A1B11AE1"))

        # Check if it's RTF disguised as .doc
        if file_bytes.startswith(b"{\\rtf"):
            for encoding in ("utf-8", "cp1252", "latin-1"):
                try:
                    text = file_bytes.decode(encoding, errors="ignore")
                    if text.strip():
                        return TextExtractionResult(text=text, extractor="rtf_in_doc")
                except Exception:
                    continue

        if not is_binary_doc:
            raise TextExtractionError(
                "Legacy .doc extraction failed. Please convert the file to .docx or PDF.",
                file_path=file_path,
            )

        # Try OCR
        if self._ocr.is_configured:
            try:
                content = await self._ocr.extract_text(file_bytes, file_path)
                if content.strip():
                    return TextExtractionResult(
                        text=content,
                        extractor=f"ocr_{self._ocr.provider_name}",
                    )
            except Exception as exc:
                logger.warning(
                    "Legacy .doc OCR extraction failed for %s: %s",
                    file_path,
                    exc,
                )

        raise TextExtractionError(
            "Legacy .doc extraction failed. Please convert the file to .docx or PDF.",
            file_path=file_path,
        )


class TextExtractionService:
    """
    Facade that routes extraction through dedicated extractor classes.

    Multi-provider OCR support - uses OCRAdapter for fallback extraction.

    Args:
        ocr_adapter: OCR adapter for scanned document fallback. If None,
            OCR fallback will be disabled (NoOpOCRAdapter).
    """

    def __init__(self, ocr_adapter: Optional[OCRAdapter] = None) -> None:
        self._ocr = ocr_adapter or NoOpOCRAdapter()
        self._extractors = (
            PlainTextExtractor(),
            PdfExtractor(self._ocr),
            DocxExtractor(),
            DocExtractor(self._ocr),
        )

    async def extract_text(self, file_bytes: bytes, file_path: str) -> TextExtractionResult:
        normalized_bytes = normalize_file_bytes(file_bytes=file_bytes, file_path=file_path)
        extension = os.path.splitext(file_path)[1].lower()

        for extractor in self._extractors:
            if extractor.can_extract(extension):
                return await extractor.extract(normalized_bytes, file_path)

        raise TextExtractionError(f"Unsupported file type: {extension}", file_path=file_path)
